import os
import sys

# Force UTF-8 stdout/stderr - Windows console defaults to cp1252 which
# can't encode chars like → (U+2192) used in our SDD templates and
# prompts. Without this, any print() containing such chars raises
# UnicodeEncodeError which propagates as a 500 error during generation.
try:
    sys.stdout.reconfigure(encoding='utf-8', errors='replace')
    sys.stderr.reconfigure(encoding='utf-8', errors='replace')
except (AttributeError, OSError):
    pass  # older Python or constrained stream

import uuid
import shutil
import json
import base64
from fastapi import FastAPI, UploadFile, File, Form, HTTPException, Body, Depends, Response, Request, WebSocket, WebSocketDisconnect
from fastapi.middleware.cors import CORSMiddleware
from typing import Optional, List
import asyncio
import websockets

from config import MAX_UPLOAD_SIZE_MB, ALL_SUPPORTED_EXTENSIONS
from graph.pipeline import pipeline
from utils.media_processor import detect_input_type
from utils.template_store import save_template, get_template_path, has_template, delete_template
from utils.template_renderer import build_template_outline, render_markdown_into_template
from pydantic import BaseModel
from utils.gemini_client import call_gemini_structured, call_gemini_text, call_gemini_multimodal, transcribe_audio_bytes, call_gemini_tts
from db.mongo_client import (
    save_assessment,
    get_all_assessments,
    get_assessment,
    delete_assessment,
    update_assessment_bpmn,
    update_assessment_chat,
    update_assessment_audio_script,
    update_assessment_document,
    revert_global_assessment,
    update_assessment_hourly_rate,
    update_assessment_data,
    save_live_chat_session,
    get_live_chat_session,
    delete_live_chat_session
)
from db.user_client import create_user, get_user_by_email, update_password, get_user_by_id
from utils.auth import get_password_hash, verify_password, create_access_token, get_current_user, create_refresh_token, verify_token, get_api_account
from db.redis_client import get_redis_client
from utils.email_sender import send_otp_email
import random
from google.oauth2 import id_token
from google.auth.transport import requests as google_requests
import config
from routes.keys import router as keys_router


app = FastAPI(title="AP Process Discovery Agent")

origins = ["http://localhost:5173", "http://127.0.0.1:5173"]
allowed_origins_env = os.getenv("ALLOWED_ORIGINS")
if allowed_origins_env:
    origins = [o.strip() for o in allowed_origins_env.split(",") if o.strip()]
elif os.getenv("FRONTEND_URL"):
    origins.append(os.getenv("FRONTEND_URL"))

app.add_middleware(
    CORSMiddleware,
    allow_origins=origins,
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

app.include_router(keys_router)

# Models for Auth
class UserSignup(BaseModel):
    name: str
    email: str
    phone: str
    password: str

class UserLogin(BaseModel):
    email: str
    password: str

class GoogleLogin(BaseModel):
    credential: str

class ForgotPasswordRequest(BaseModel):
    email: str

class VerifyOTPRequest(BaseModel):
    email: str
    otp: str

class ResetPasswordRequest(BaseModel):
    reset_token: str
    new_password: str

@app.post("/api/auth/signup")
async def signup(user: UserSignup):
    """Register a new user."""
    # 1. Check if user exists
    existing_user = get_user_by_email(user.email)
    if existing_user:
        raise HTTPException(status_code=400, detail="Email already registered")
        
    # 2. Hash password
    hashed_pw = get_password_hash(user.password)
    
    # 3. Save to DB
    user_data = {
        "name": user.name,
        "email": user.email,
        "phone": user.phone,
        "password": hashed_pw,
        "auth_provider": "local"
    }
    try:
        create_user(user_data)
        return {"status": "success", "message": "User created successfully"}
    except Exception as e:
        raise HTTPException(status_code=500, detail="Database error: " + str(e))
def set_auth_cookies(response: Response, request: Request, access_token: str, refresh_token: str):
    hostname = request.url.hostname or ""
    is_local = "localhost" in hostname or "127.0.0.1" in hostname
    secure = not is_local
    samesite = "none" if not is_local else "lax"
    
    response.set_cookie(key="access_token", value=access_token, httponly=True, secure=secure, samesite=samesite, max_age=15*60)
    response.set_cookie(key="refresh_token", value=refresh_token, httponly=True, secure=secure, samesite=samesite, max_age=7*24*60*60)

@app.post("/api/auth/login")
async def login(user: UserLogin, response: Response, request: Request):
    """Authenticate a user and set JWT cookies."""
    # 1. Find user
    db_user = get_user_by_email(user.email)
    if not db_user:
        raise HTTPException(status_code=401, detail="Invalid email or password")
        
    if db_user.get("auth_provider") == "google":
        raise HTTPException(status_code=400, detail="You signed up with Google before. Please use 'Continue with Google' to log in.")
        
    # 2. Verify password
    if not verify_password(user.password, db_user["password"]):
        raise HTTPException(status_code=401, detail="Invalid email or password")
        
    # 3. Generate tokens
    access_token = create_access_token(data={"sub": str(db_user["_id"])})
    refresh_token = create_refresh_token(data={"sub": str(db_user["_id"])})
    
    # 4. Set cookies — dynamic secure/samesite for production/dev
    set_auth_cookies(response, request, access_token, refresh_token)
    
    return {"status": "success", "access_token": access_token}

@app.get("/api/auth/check")
async def check_auth(request: Request, user_id: str = Depends(get_current_user)):
    """Check if the current session is valid."""
    user = get_user_by_id(user_id)
    if not user:
        raise HTTPException(status_code=401, detail="User not found")
    token = request.cookies.get("access_token")
    return {"status": "success", "user": {"id": user_id, "name": user["name"], "email": user["email"]}, "access_token": token}

@app.post("/api/auth/refresh")
async def refresh_token_endpoint(request: Request, response: Response):
    """Refresh the access token using a valid refresh token."""
    refresh_token = request.cookies.get("refresh_token")
    if not refresh_token:
        raise HTTPException(status_code=401, detail="No refresh token")
    try:
        payload = verify_token(refresh_token)
        if payload.get("type") != "refresh":
            raise HTTPException(status_code=401, detail="Invalid token type")
        user_id = payload.get("sub")
        
        # issue new access token
        new_access = create_access_token(data={"sub": user_id})
        
        # Dynamic cookie options
        hostname = request.url.hostname or ""
        is_local = "localhost" in hostname or "127.0.0.1" in hostname
        secure = not is_local
        samesite = "none" if not is_local else "lax"
        
        response.set_cookie(key="access_token", value=new_access, httponly=True, secure=secure, samesite=samesite, max_age=15*60)
        return {"status": "success", "access_token": new_access}
    except Exception:
        raise HTTPException(status_code=401, detail="Invalid or expired refresh token")

@app.post("/api/auth/logout")
async def logout(request: Request, response: Response):
    """Log out by clearing HTTP-only cookies."""
    hostname = request.url.hostname or ""
    is_local = "localhost" in hostname or "127.0.0.1" in hostname
    secure = not is_local
    samesite = "none" if not is_local else "lax"
    
    response.delete_cookie("access_token", secure=secure, samesite=samesite)
    response.delete_cookie("refresh_token", secure=secure, samesite=samesite)
    return {"status": "success"}

@app.post("/api/auth/google")
async def google_login(data: GoogleLogin, response: Response, request: Request):
    """Authenticate a user using Google OAuth credential."""
    try:
        # Verify the token
        idinfo = id_token.verify_oauth2_token(
            data.credential, 
            google_requests.Request(), 
            config.GOOGLE_CLIENT_ID
        )
        
        email = idinfo.get("email")
        name = idinfo.get("name")
        
        if not email:
            raise HTTPException(status_code=400, detail="Google token did not contain an email")
            
        # Check if user exists
        db_user = get_user_by_email(email)
        
        if not db_user:
            # Create a new user with a random unguessable password
            random_password = str(uuid.uuid4())
            hashed_pw = get_password_hash(random_password)
            user_id_str = create_user({
                "name": name or email.split("@")[0],
                "email": email,
                "password": hashed_pw,
                "auth_provider": "google"
            })
        else:
            user_id_str = str(db_user["_id"])
            
        # Generate tokens
        access_token = create_access_token(data={"sub": user_id_str})
        refresh_token = create_refresh_token(data={"sub": user_id_str})
        
        # Set cookies — dynamic secure/samesite for production/dev
        set_auth_cookies(response, request, access_token, refresh_token)
        
        return {"status": "success", "access_token": access_token}
    except HTTPException:
        raise
    except ValueError as e:
        print(f"[Google Auth] Token verification failed (ValueError): {e}")
        raise HTTPException(status_code=401, detail="Invalid Google token")
    except Exception as e:
        print(f"[Google Auth] Unexpected error: {type(e).__name__}: {e}")
        raise HTTPException(status_code=500, detail=f"Google login failed: {str(e)}")

@app.post("/api/auth/forgot-password")
async def forgot_password(req: ForgotPasswordRequest):
    """Generate OTP and send it via email."""
    user = get_user_by_email(req.email)
    if not user:
        # We still return success to prevent email enumeration, but we don't send anything
        return {"status": "success", "message": "If the email is registered, an OTP will be sent."}
        
    if user.get("auth_provider") == "google":
        raise HTTPException(status_code=400, detail="Your account is linked to Google. Please use 'Continue with Google' to log in.")
        
    redis = get_redis_client()
    if not redis:
        raise HTTPException(status_code=500, detail="Redis connection failed")
        
    # Generate 6-digit OTP
    otp = str(random.randint(100000, 999999))
    
    # Store in Redis with 5 minutes expiry (300 seconds)
    redis.setex(f"otp:{req.email}", 300, otp)
    
    # Send email
    success = send_otp_email(req.email, otp)
    if not success:
        raise HTTPException(status_code=500, detail="Failed to send OTP email")
        
    return {"status": "success", "message": "If the email is registered, an OTP will be sent."}

@app.post("/api/auth/verify-otp")
async def verify_otp(req: VerifyOTPRequest):
    """Verify OTP and return a reset token."""
    redis = get_redis_client()
    if not redis:
        raise HTTPException(status_code=500, detail="Redis connection failed")
        
    stored_otp = redis.get(f"otp:{req.email}")
    if not stored_otp or stored_otp != req.otp:
        raise HTTPException(status_code=400, detail="Invalid or expired OTP")
        
    # Valid OTP: Generate a secure reset token
    reset_token = str(uuid.uuid4())
    
    # Store the token mapped to the email (15 minutes expiry)
    redis.setex(f"reset_token:{reset_token}", 900, req.email)
    
    # Delete the used OTP
    redis.delete(f"otp:{req.email}")
    
    return {"status": "success", "reset_token": reset_token}

@app.post("/api/auth/reset-password")
async def reset_password(req: ResetPasswordRequest):
    """Reset the password using the reset token."""
    redis = get_redis_client()
    if not redis:
        raise HTTPException(status_code=500, detail="Redis connection failed")
        
    email = redis.get(f"reset_token:{req.reset_token}")
    if not email:
        raise HTTPException(status_code=400, detail="Invalid or expired reset token")
        
    # Hash the new password
    hashed_pw = get_password_hash(req.new_password)
    
    # Update in DB
    success = update_password(email, hashed_pw)
    if not success:
        raise HTTPException(status_code=500, detail="Failed to update password")
        
    # Delete the reset token so it can't be reused
    redis.delete(f"reset_token:{req.reset_token}")
    
    return {"status": "success", "message": "Password updated successfully"}



TEMP_DIR = os.path.join(os.path.dirname(__file__), "temp")
os.makedirs(TEMP_DIR, exist_ok=True)


@app.get("/api/health")
async def health_check():
    return {"status": "ok", "message": "AP Discovery Agent is running"}


@app.post("/api/transcribe")
async def transcribe_audio(
    file: UploadFile = File(...),
    user_id: str = Depends(get_current_user)
):
    """Transcribe an audio recording to text using Gemini 2.5 Flash."""
    try:
        # Save the uploaded audio to a temp file
        ext = os.path.splitext(file.filename or "recording.webm")[1] or ".webm"
        temp_filename = f"{uuid.uuid4()}{ext}"
        temp_path = os.path.join(TEMP_DIR, temp_filename)

        contents = await file.read()
        if len(contents) < 1000:  # Less than ~1KB is likely too short / empty
            raise HTTPException(status_code=400, detail="Recording is too short. Please speak for at least a second.")

        with open(temp_path, "wb") as f:
            f.write(contents)

        # Use Gemini multimodal to transcribe
        transcription_prompt = (
            "You are an expert multilingual transcriber. "
            "Transcribe the following audio recording faithfully and accurately. "
            "Preserve the original language(s) spoken (e.g. English, Hindi, Tamil, etc.). "
            "If the speaker mixes languages (code-switching), transcribe each part in its original language. "
            "Output ONLY the plain transcription text, nothing else — no timestamps, no labels, no formatting."
        )

        # Detect mime type from extension
        mime_map = {
            '.webm': 'audio/webm',
            '.wav': 'audio/wav',
            '.mp3': 'audio/mpeg',
            '.m4a': 'audio/mp4',
            '.ogg': 'audio/ogg',
        }
        mime_type = mime_map.get(ext, 'audio/webm')

        # Use inline bytes transcription (avoids File API processing failures on short clips)
        transcribed_text = transcribe_audio_bytes(
            file_path=temp_path,
            prompt=transcription_prompt,
            mime_type=mime_type,
            model=config.GEMINI_MODEL_FAST  # gemini-2.5-flash — best for audio
        )

        # Cleanup temp file
        try:
            os.remove(temp_path)
        except Exception:
            pass

        # Handle empty/None response
        if not transcribed_text:
            return {"status": "success", "data": {"text": "(Could not transcribe audio. Please try again.)"}}

        return {"status": "success", "data": {"text": transcribed_text.strip()}}

    except HTTPException:
        raise
    except Exception as e:
        # Cleanup on error
        try:
            os.remove(temp_path)
        except Exception:
            pass
        raise HTTPException(status_code=500, detail=f"Transcription error: {str(e)}")


@app.post("/api/analyze")
async def analyze(
    text: Optional[str] = Form(None),
    file: Optional[UploadFile] = File(None),
    followup_context: Optional[str] = Form(None),
    user_id: str = Depends(get_current_user)
):
    """Analyze an AP process description from text or file upload."""
    try:
        if file:
            # Validate file extension
            ext = os.path.splitext(file.filename)[1].lower()
            if ext not in ALL_SUPPORTED_EXTENSIONS:
                raise HTTPException(
                    status_code=415,
                    detail=f"Unsupported file type: {ext}. Supported: {ALL_SUPPORTED_EXTENSIONS}"
                )

            # Check file size
            contents = await file.read()
            size_mb = len(contents) / (1024 * 1024)
            if size_mb > MAX_UPLOAD_SIZE_MB:
                raise HTTPException(
                    status_code=413,
                    detail=f"File too large: {size_mb:.1f}MB. Max: {MAX_UPLOAD_SIZE_MB}MB"
                )

            # Save to temp directory
            temp_filename = f"{uuid.uuid4()}{ext}"
            temp_path = os.path.join(TEMP_DIR, temp_filename)
            with open(temp_path, "wb") as f:
                f.write(contents)

            input_format = detect_input_type(file.filename)

            # Run pipeline
            result = pipeline.invoke({
                "raw_text": "",
                "input_format": input_format,
                "original_filename": file.filename,
                "file_path": temp_path,
            })

            # Cleanup temp file
            try:
                os.remove(temp_path)
            except Exception:
                pass

        elif text:
            if len(text.strip()) < 50:
                raise HTTPException(
                    status_code=400,
                    detail="Text too short. Please provide at least 50 characters."
                )

            # Enrich text with follow-up Q&A context if provided
            enriched_text = text
            if followup_context and followup_context.strip():
                enriched_text = followup_context.strip() + "\n\n---\nOriginal Transcript:\n" + text

            result = pipeline.invoke({
                "raw_text": enriched_text,
                "input_format": "text",
                "original_filename": "",
                "file_path": "",
            })

        else:
            raise HTTPException(
                status_code=400,
                detail="Please provide either text or a file."
            )

        # Build assessment data structure
        assessment_data = {
            "original_filename": result.get("original_filename", ""),
            "input_format": result.get("input_format", "text"),
            "process_map": result.get("process_map", []),
            "scored_steps": result.get("scored_steps", []),
            "priority_targets": result.get("priority_targets", []),
            "opportunities": result.get("opportunities", []),
            "executive_summary": result.get("executive_summary", ""),
            "markdown_report": result.get("markdown_report", ""),
            "systems_mentioned": result.get("systems_mentioned", []),
            "roles_identified": result.get("roles_identified", []),
            "pain_points": result.get("pain_points", []),
            "hourly_rate": result.get("hourly_rate"),
            "original_transcript": enriched_text if text else "",
            "project_timeline": result.get("project_timeline", {}),
            "roi_estimate": result.get("roi_estimate", {}),
            "discovery_facts": result.get("discovery_facts", {}),
            "currency": result.get("currency"),
            "missing_critical_facts": result.get("missing_critical_facts", []),
        }

        # Save to database
        doc_id = save_assessment(assessment_data, user_id)
        if doc_id:
            assessment_data["id"] = doc_id

        # Return the assessment result
        return {
            "status": "success",
            "data": assessment_data,
        }

    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Pipeline error: {str(e)}")


# ── Interactive Analysis Mode — Follow-Up Questions ────────────────────────────

class FollowUpQuestionsRequest(BaseModel):
    text: Optional[str] = None
    transcript: Optional[str] = None  # Pre-extracted transcript (for file uploads handled on frontend)

class FollowUpAnswersRequest(BaseModel):
    transcript: str
    conversation: list  # List of {"question": str, "answer": str} dicts


@app.post("/api/analyze/followup-questions")
async def generate_followup_questions(
    text: Optional[str] = Form(None),
    file: Optional[UploadFile] = File(None),
    user_id: str = Depends(get_current_user)
):
    """Read the transcript and generate follow-up questions for interactive analysis."""
    try:
        transcript = ""

        if file:
            # Validate file extension
            ext = os.path.splitext(file.filename)[1].lower()
            if ext not in ALL_SUPPORTED_EXTENSIONS:
                raise HTTPException(
                    status_code=415,
                    detail=f"Unsupported file type: {ext}. Supported: {ALL_SUPPORTED_EXTENSIONS}"
                )

            # Check file size
            contents = await file.read()
            size_mb = len(contents) / (1024 * 1024)
            if size_mb > MAX_UPLOAD_SIZE_MB:
                raise HTTPException(
                    status_code=413,
                    detail=f"File too large: {size_mb:.1f}MB. Max: {MAX_UPLOAD_SIZE_MB}MB"
                )

            # Save to temp directory
            temp_filename = f"{uuid.uuid4()}{ext}"
            temp_path = os.path.join(TEMP_DIR, temp_filename)
            with open(temp_path, "wb") as f:
                f.write(contents)

            # Extract text from the file using the same logic as the pipeline
            input_format = detect_input_type(file.filename)
            if input_format == "text":
                transcript = contents.decode("utf-8", errors="replace")
            elif input_format == "docx":
                from utils.docx_parser import extract_text_from_docx
                transcript = extract_text_from_docx(temp_path)
            elif input_format == "audio":
                from utils.media_processor import process_audio
                transcript = process_audio(temp_path)
            elif input_format == "video":
                from utils.media_processor import process_video
                transcript = process_video(temp_path)

            # Cleanup temp file
            try:
                os.remove(temp_path)
            except Exception:
                pass

        elif text:
            if len(text.strip()) < 50:
                raise HTTPException(
                    status_code=400,
                    detail="Text too short. Please provide at least 50 characters."
                )
            transcript = text
        else:
            raise HTTPException(
                status_code=400,
                detail="Please provide either text or a file."
            )

        # Generate follow-up questions using Gemini
        prompt = f"""You are an expert Accounts Payable (AP) automation analyst conducting a discovery session.

You have just received the following transcript describing an AP process. Read it carefully and identify areas that are:
- Unclear or ambiguous (e.g., "we send it somewhere" — where exactly?)
- Missing critical details (e.g., no mention of approval thresholds, volumes, or frequencies)
- Lacking system/technology specifics (e.g., which ERP, what tool for matching?)
- Missing exception handling details (e.g., what happens when an invoice doesn't match?)
- Unclear role assignments (e.g., who approves, who reviews?)
- Missing cost information (e.g., hourly rate of AP staff, cost per invoice, team size)

IMPORTANT RULE: The downstream ROI is calculated deterministically from FIVE critical facts. For ANY of these that the transcript does not already state clearly, you MUST include a specific question to capture it:
  1. Currency (USD / GBP / EUR / etc.)
  2. Loaded hourly rate (or salary band) of the AP staff
  3. Transaction/invoice volume that can be annualized
  4. Total FTE / headcount working on this process
  5. Exception rate — the % of transactions needing manual handling
Always include the cost/hourly-rate question. Only skip one of the five if the transcript already answers it clearly.

Generate 4-6 targeted follow-up questions that will help produce a significantly more accurate and detailed automation analysis, prioritizing any of the five critical facts that are missing.

IMPORTANT: Output ONLY a JSON array of question strings. No markdown, no explanation.
Example: ["Which currency should the ROI be calculated in (e.g. USD, GBP, EUR)?", "What is the approximate loaded hourly rate of the AP team members involved?", "Roughly how many invoices are processed per month or year?", "About what percentage of invoices hit an exception or need manual handling?"]

Transcript:
{transcript}
"""

        questions_raw = call_gemini_text(
            prompt=prompt,
            system_prompt="You are an AP automation discovery analyst. Output ONLY a valid JSON array of question strings.",
            model="gemini-3.5-flash"
        )

        # Parse the JSON array
        import re
        cleaned = questions_raw.strip()
        if cleaned.startswith("```json"):
            cleaned = cleaned[7:]
        elif cleaned.startswith("```"):
            cleaned = cleaned[3:]
        if cleaned.endswith("```"):
            cleaned = cleaned[:-3]
        cleaned = cleaned.strip()

        try:
            questions = json.loads(cleaned)
        except json.JSONDecodeError:
            # Fallback: try to extract array from the text
            match = re.search(r'\[.*\]', cleaned, re.DOTALL)
            if match:
                questions = json.loads(match.group())
            else:
                questions = ["Could you describe the end-to-end AP process in more detail?",
                            "What systems and tools are used in each step?",
                            "What are the common exceptions or issues you encounter?"]

        return {
            "status": "success",
            "data": {
                "questions": questions,
                "transcript": transcript
            }
        }

    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Follow-up generation error: {str(e)}")


@app.post("/api/analyze/followup-answers")
async def process_followup_answers(
    payload: FollowUpAnswersRequest,
    user_id: str = Depends(get_current_user)
):
    """Process user answers and decide if more questions are needed."""
    try:
        transcript = payload.transcript
        conversation = payload.conversation

        # Format the existing Q&A
        qa_str = ""
        for qa in conversation:
            qa_str += f"Q: {qa.get('question', '')}\nA: {qa.get('answer', '')}\n\n"

        prompt = f"""You are an expert AP automation analyst conducting a follow-up discovery session.

You previously asked questions about an AP process transcript, and the user has answered them.
Review the transcript and the Q&A to decide if you need any MORE clarification.

Transcript:
{transcript}

---
Previous Q&A:
{qa_str}
---

If the answers are sufficient for a thorough automation analysis, respond with:
{{"satisfied": true, "questions": []}}

If you still need more clarity, respond with 2-3 additional questions:
{{"satisfied": false, "questions": ["question1", "question2"]}}

IMPORTANT: Output ONLY a valid JSON object. No markdown, no explanation.
"""

        result_raw = call_gemini_text(
            prompt=prompt,
            system_prompt="You are an AP automation discovery analyst. Output ONLY a valid JSON object.",
            model="gemini-3.5-flash"
        )

        # Parse the JSON
        cleaned = result_raw.strip()
        if cleaned.startswith("```json"):
            cleaned = cleaned[7:]
        elif cleaned.startswith("```"):
            cleaned = cleaned[3:]
        if cleaned.endswith("```"):
            cleaned = cleaned[:-3]
        cleaned = cleaned.strip()

        try:
            result = json.loads(cleaned)
        except json.JSONDecodeError:
            result = {"satisfied": True, "questions": []}

        return {
            "status": "success",
            "data": result
        }

    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Follow-up processing error: {str(e)}")


@app.get("/api/history")
async def list_history(user_id: str = Depends(get_current_user)):
    """Get a list of past assessments."""
    try:
        assessments = get_all_assessments(user_id)
        return {"status": "success", "data": assessments}
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


@app.get("/api/history/{assessment_id}")
async def get_history_item(assessment_id: str, user_id: str = Depends(get_current_user)):
    """Retrieve a full past assessment."""
    try:
        assessment = get_assessment(assessment_id, user_id)
        if not assessment:
            raise HTTPException(status_code=404, detail="Assessment not found")
        return {"status": "success", "data": assessment}
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


@app.delete("/api/history/{assessment_id}")
async def delete_history_item(assessment_id: str, user_id: str = Depends(get_current_user)):
    """Delete a past assessment."""
    try:
        success = delete_assessment(assessment_id, user_id)
        if not success:
            raise HTTPException(status_code=404, detail="Assessment not found or could not be deleted")
        return {"status": "success", "message": "Assessment deleted successfully"}
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


class BpmnUpdateRequest(BaseModel):
    bpmn_xml: str

@app.put("/api/history/{assessment_id}/bpmn")
async def update_bpmn_diagram(assessment_id: str, payload: BpmnUpdateRequest, user_id: str = Depends(get_current_user)):
    """Update the BPMN diagram XML of a past assessment."""
    try:
        success = update_assessment_bpmn(assessment_id, user_id, payload.bpmn_xml)
        if not success:
            raise HTTPException(status_code=404, detail="Assessment not found or could not be updated")
        return {"status": "success", "message": "BPMN diagram updated successfully"}
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


class ChatMessage(BaseModel):
    role: str
    content: str

class ChatHistoryUpdateRequest(BaseModel):
    chat_history: List[ChatMessage]

@app.put("/api/history/{assessment_id}/chat")
async def update_chat_history(assessment_id: str, payload: ChatHistoryUpdateRequest, user_id: str = Depends(get_current_user)):
    """Update the chat history of a past assessment."""
    try:
        # Convert Pydantic models back to dictionaries
        history_list = [msg.model_dump() if hasattr(msg, 'model_dump') else msg.dict() for msg in payload.chat_history]
        success = update_assessment_chat(assessment_id, user_id, history_list)
        if not success:
            raise HTTPException(status_code=404, detail="Assessment not found or could not be updated")
        return {"status": "success", "message": "Chat history updated successfully"}
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

class HourlyRateUpdateRequest(BaseModel):
    hourly_rate: float

@app.put("/api/history/{assessment_id}/hourly-rate")
async def update_hourly_rate(assessment_id: str, payload: HourlyRateUpdateRequest, user_id: str = Depends(get_current_user)):
    """Update the hourly rate of a past assessment."""
    try:
        success = update_assessment_hourly_rate(assessment_id, user_id, payload.hourly_rate)
        if not success:
            raise HTTPException(status_code=404, detail="Assessment not found or could not be updated")
        return {"status": "success", "message": "Hourly rate updated successfully"}
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

class AudioOverviewRequest(BaseModel):
    language: str = "English"

@app.post("/api/history/{assessment_id}/audio-overview")
async def create_audio_overview(assessment_id: str, payload: AudioOverviewRequest, user_id: str = Depends(get_current_user)):
    """Generate a podcast-style audio overview of an assessment using Gemini TTS."""
    try:
        # Get assessment
        assessment = get_assessment(assessment_id, user_id)
        if not assessment:
            raise HTTPException(status_code=404, detail="Assessment not found")

        # Remove massive fields like bpmn_xml and chat_history
        context_data = {k: v for k, v in assessment.items() if k not in ('bpmn_xml', 'chat_history', '_id', 'user_id')}

        # Check if the script is already saved in MongoDB
        audio_scripts = assessment.get("audio_scripts", {})
        script = audio_scripts.get(payload.language)

        if not script:
            context_str = json.dumps(context_data, indent=2)
            prompt = f"""You are an expert AP Automation podcast host recording a detailed audio overview.

Write a SPOKEN SCRIPT of approximately 650 words (about 4.5 minutes when spoken aloud).
Do NOT use markdown, bullet points, headers, asterisks, or any formatting — plain natural spoken text only.
Do NOT include stage directions, speaker labels, or anything that isn't meant to be spoken.

Cover ALL of these topics in your script:
1. A warm greeting and introduction to what this podcast episode is about
2. The current AP process — how many steps, who is involved, what the daily workflow looks like
3. The key bottlenecks and pain points you discovered
4. The Automation Complexity Score and what it means for this process
5. The specific automation opportunities and priority targets identified
6. What the automated solution will look like — the bot's workflow
7. The projected business impact: time saved, cost reduction, error reduction
8. A closing note encouraging the team to move forward with automation

Tone: warm, authoritative, conversational, enthusiastic — like a real analyst presenting findings live.
Language: {payload.language} (speak entirely and fluently in this language only).

Analysis data:
{context_str[:6000]}
"""
            # Step 1: Generate the podcast script using Gemini
            script = call_gemini_text(
                prompt=prompt,
                system_prompt="You are a professional podcast host. Write exactly 650 words of plain spoken text with no formatting."
            )

            # Save the script to MongoDB for future use
            update_assessment_audio_script(assessment_id, user_id, payload.language, script)

        # Step 2: Convert script to audio using chunked Gemini TTS
        # call_gemini_tts automatically splits long text into ~2800-char chunks,
        # calls TTS for each chunk, and concatenates the raw PCM bytes.
        voice_map = {
            'Hindi':           'Aoede',
            'Punjabi':         'Aoede',
            'Tamil':           'Aoede',
            'Telugu':          'Aoede',
            'Bengali':         'Aoede',
            'Marathi':         'Aoede',
            'Gujarati':        'Aoede',
            'Malayalam':       'Aoede',
            'Kannada':         'Aoede',
            'Korean':          'Leda',
            'Japanese':        'Leda',
            'Mandarin Chinese':'Leda',
            'Russian':         'Charon',
            'Spanish':         'Zephyr',
            'French':          'Zephyr',
            'German':          'Charon',
            'English':         'Kore',
        }
        voice_name = voice_map.get(payload.language, 'Kore')
        print(f"[TTS] Script length: {len(script)} chars | Voice: {voice_name} | Lang: {payload.language}")

        pcm_bytes = call_gemini_tts(text=script, voice_name=voice_name)

        if not pcm_bytes:
            raise HTTPException(status_code=500, detail="Gemini TTS returned no audio data.")

        # Wrap raw PCM (24kHz, 16-bit, mono) into a valid WAV file
        import wave
        import io
        wav_buffer = io.BytesIO()
        with wave.open(wav_buffer, 'wb') as wf:
            wf.setnchannels(1)       # mono
            wf.setsampwidth(2)       # 16-bit = 2 bytes per sample
            wf.setframerate(24000)   # 24kHz — Gemini TTS sample rate
            wf.writeframes(pcm_bytes)
        wav_buffer.seek(0)
        audio_bytes = wav_buffer.read()

        return Response(content=audio_bytes, media_type="audio/wav")
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


# ── PDD / SDD Document Generation ──────────────────────────────────────────────

PDD_TEMPLATE = """You are a Process Design Document (PDD) specialist agent for RPA automation projects. Your job is to produce a complete, professionally formatted PDD that follows the standard UiPath PDD template structure, based on the AP assessment data provided below.

CRITICAL RULES FOR DIAGRAMS:
- NEVER use Mermaid.js, ASCII art, box-drawing characters, or plain-text diagrams anywhere in this document.
- For any diagram (architecture, flowchart, decision tree, etc.) EXCEPT the main BPMN process map, output a placeholder in EXACTLY this format:
  [[FLOWCHART: <a precise, specific description of what this diagram should visualise — include step names, decision branches, system names, and flow direction>]]
- The placeholder will be replaced by a real AI-generated image at render time. Write the description inside the placeholder as if instructing a visual designer.
- For the main process flowchart (Section II.3), follow the placeholder instruction there exactly — use {{BPMN_DIAGRAM}} only.
- Example of a valid flowchart placeholder:
  [[FLOWCHART: End-to-end AP invoice processing flow. Left to right. Start: Email arrives in shared mailbox. Step 1: Bot downloads attachment. Step 2: OCR extracts invoice data. Decision: 3-way match passes? Yes → post to SAP FB60. No → flag exception and email vendor. End: Payment queued in SAP F110.]]

OUTPUT FORMAT RULES:
- Use Markdown formatting — # for section headers, ## for subsections, tables for all tabular data.
- Preserve the EXACT section numbering from the template (I.1, I.2, II.1, II.2, etc.).
- Use n/a for any fields that don't apply — never leave cells blank.
- Populate standard boilerplate (like the Purpose statement in I.1) — adapt it to the specific process name and org.
- Flag incomplete sections clearly with [⚠ TO BE COMPLETED BY SME] so the reviewer knows what still needs input.
- Tables must have headers and be properly formatted.
- CRITICAL — NEVER use HTML tags like <br>, <br/>, or &lt;br&gt; inside table cells. Markdown table cells do NOT support HTML line breaks. Use a semicolon (;) to separate multiple values in a single cell instead.
- Process steps in II.4 should be numbered sequentially and include as much detail as the data provides.
- The document should read as a professional, review-ready artifact — not a rough draft.

Generate the COMPLETE PDD using the following EXACT structure:

---

# [Process Name] — Process Design Document (PDD)
## RPA Automation Project | UiPath

---

## Cover & Metadata

**Document Title:** [Process Name] — Process Design Document
**Project:** AP Automation with UiPath RPA

| Date | Version | Role | Name | Organization/Dept | Function | Comments |
|------|---------|------|------|-------------------|----------|----------|
| [today's date] | 1.0 | AP Discovery Agent | AI-Generated | [Org from data] | Document Author | Initial draft |

**Document Approval Flow:**

| Version | Flow Stage | Role | Name | Organization | Signature/Date |
|---------|-----------|------|------|--------------|----------------|
| 1.0 | Draft Review | Process SME | [⚠ TO BE COMPLETED BY SME] | [⚠ TO BE COMPLETED BY SME] | [⚠ TO BE COMPLETED BY SME] |
| 1.0 | Approval | Process Owner | [⚠ TO BE COMPLETED BY SME] | [⚠ TO BE COMPLETED BY SME] | [⚠ TO BE COMPLETED BY SME] |

---

## I. Introduction

### I.1 Purpose of the Document
This Process Design Document (PDD) describes the as-is and to-be process for the [Process Name] automation initiative within [Organization/Dept]. It provides the detailed process understanding necessary for the RPA development team to design, build, and test the UiPath automation solution. This document serves as the primary reference for all process-related decisions throughout the automation lifecycle.

### I.2 Objectives
Describe the business objectives and quantified benefits expected after automation. Use the assessment data to populate specific ROI estimates, FTE savings, volume targets, and accuracy improvements.

### I.3 Key Contacts

| Role | Name | Email / Contact | Department |
|------|------|-----------------|------------|
| Process SME | [⚠ TO BE COMPLETED BY SME] | [⚠ TO BE COMPLETED BY SME] | [from data] |
| Process Reviewer | [⚠ TO BE COMPLETED BY SME] | [⚠ TO BE COMPLETED BY SME] | [⚠ TO BE COMPLETED BY SME] |
| Process Owner / Approver | [⚠ TO BE COMPLETED BY SME] | [⚠ TO BE COMPLETED BY SME] | [⚠ TO BE COMPLETED BY SME] |

### I.4 Minimum Prerequisites for Automation

- [ ] Completed and approved PDD
- [ ] Test data available for all exception scenarios
- [ ] User access and licenses provisioned for bot
- [ ] Credentials for all target systems secured
- [ ] All dependent systems and APIs accessible in test environment

---

## II. As-Is Process Description

### II.1 Process Overview

| Field | Details |
|-------|---------|
| Process Full Name | [from assessment data] |
| Business Area | [from assessment data] |
| Department | [from assessment data] |
| Short Description | [from assessment data] |
| Roles Involved | [from roles_identified] |
| Schedule / Frequency | [from assessment data] |
| # Items per Reference Period | [from assessment data] |
| Average Handling Time (per item) | [from assessment data] |
| Peak Periods | [from assessment data or n/a] |
| Peak Volume | [from assessment data or n/a] |
| Total FTEs | [from assessment data] |
| Expected Volume Increase | [from assessment data or n/a] |
| Exception Rate Level | [Low / Medium / High based on pain_points] |
| Input Data | [from assessment data] |
| Output Data | [from assessment data] |

### II.2 Applications Used in the Process

| Application Name & Version | System Language | Thin/Thick Client | Environment / Access Method | Comments |
|---------------------------|----------------|-------------------|----------------------------|----------|
[Populate from systems_mentioned in assessment data. Use n/a for unknown fields.]

### II.3 As-Is Process Map

Instead of generating a Mermaid flowchart here, output exactly this placeholder string: {{BPMN_DIAGRAM}}

### II.4 Detailed As-Is Process Steps

| Step # | Input | Description | Details (Screen/Doc Ref) | Exception Handling | Possible Actions | Business Rules |
|--------|-------|-------------|--------------------------|-------------------|-----------------|----------------|
[Populate one row per step from process_map / scored_steps. Number sequentially. Include as much detail as the data provides.]

### II.5 Input Data Description

| Step | Sample Reference | Input Type | Location | Inputs Standard? (Y/N) | Inputs Structured? (Y/N) | Data Fields Used |
|------|-----------------|------------|----------|------------------------|--------------------------|-----------------|
[Populate from the process steps in the assessment data.]

---

## III. To-Be Process Description

### III.1 To-Be Detailed Process Map

Describe which steps will be handled by the UiPath bot vs. which remain with humans. Reference the priority_targets and opportunities from the assessment data. Note the automation triggers, outputs, and handoff points.

### III.2 Parallel Initiatives / Overlap

| Initiative | Impact on This Project | Expected Completion | Contact |
|-----------|----------------------|--------------------|---------| 
| [⚠ TO BE COMPLETED BY SME] | [⚠ TO BE COMPLETED BY SME] | [⚠ TO BE COMPLETED BY SME] | [⚠ TO BE COMPLETED BY SME] |

### III.3 In Scope for RPA

Based on the opportunities and priority_targets from the assessment data, list the specific steps and activities the robot will handle:

[Use bullet points to list each in-scope activity derived from the assessment]

### III.4 Out of Scope for RPA

| Sub-Process | Activity / Step | Reason for Exclusion | Impact on To-Be | Future Considerations |
|------------|-----------------|---------------------|----------------|-----------------------|
[Populate based on pain_points and steps with low ACS scores. Use n/a if none identified.]

### III.5 Business Exceptions Handling

**Known Business Exceptions:**

| BE# | Exception Name | Triggered at Step | Parameters / Conditions | Action |
|-----|---------------|-------------------|------------------------|--------|
[Populate from pain_points and exception handling in the assessment. Number each BE.]

**Unknown Business Exceptions:** Any business exception not covered in the table above will be flagged and routed to the human operator queue for manual review. The bot will log the exception with full context (transaction ID, step, timestamp) and send a notification to the Process SME.

### III.6 Application Error and Exception Handling

**Known Application Errors:**

| # | Error Name | Triggered at Step | Parameters / Conditions | Action |
|---|-----------|-------------------|------------------------|--------|
[Populate from system errors / technical risks identified in the assessment. Use n/a if none documented.]

**Unknown Application Errors:** For any unhandled system error, the bot will take a screenshot, log the error details, stop processing the current transaction, and alert the support team via the configured notification channel.

### III.7 Reporting

| # | Report Type | Update Frequency | Details | Monitoring Tool |
|---|------------|-----------------|---------|-----------------|
| 1 | Transaction Log | Real-time | Log of all processed items with status (Success / Exception / Failed) | UiPath Orchestrator |
| 2 | Daily Summary | Daily | Count of items processed, exceptions flagged, FTE hours saved | UiPath Orchestrator / Email |
| 3 | Exception Report | On-demand | List of all business and system exceptions with resolution status | UiPath Orchestrator |

---

## IV. Return on Investment (ROI) Projection

Based on the automation potential, provide a brief summary of the projected Return on Investment. 
Focus on:
- **Estimated Effort Reduction:** [Calculate from opportunities]
- **Time/Cost Savings Projection:** How much manual time and relative cost will be saved per cycle.
- **Estimated Development Effort:** [Calculate total estimated_development_days from opportunities] days (Includes Development, Testing, and Hypercare phases).
- **Payback Timeline:** A high-level estimate of the break-even period based on standard automation costs vs the projected savings.

---

## V. Other Observations

[Include any additional relevant observations from the assessment data, such as audit requirements, compliance considerations, seasonal volume spikes, specific monitoring needs, or data quality concerns. If none, write: No additional observations at this stage.]

---

## VI. Additional Sources of Process Documentation

| Document Type | Title / Description | Link / Location | Comments |
|--------------|--------------------|-----------------|---------| 
| Assessment Data | AI-generated AP process analysis | Internal system | Source data used to generate this PDD |
| [⚠ TO BE COMPLETED BY SME] | [Video recordings, SOPs, business logic tables] | [⚠ TO BE COMPLETED BY SME] | [⚠ TO BE COMPLETED BY SME] |

---

Here is the assessment data to base this PDD on:
{context}
"""

SDD_TEMPLATE = """You are a Solution Design Document (SDD) specialist agent for RPA automation projects. Your job is to produce a complete, professionally formatted SDD that follows the standard UiPath SDD 10-sheet template structure, based on the AP assessment data provided below.

OUTPUT FORMAT RULES:
- Structure output by sheet — use a clear --- divider and sheet header (## Sheet N: Name) for each section.
- All tables must have headers and be formatted in Markdown table syntax.
- CRITICAL — NEVER use HTML tags like <br>, <br/>, or &lt;br&gt; inside table cells. Markdown table cells do NOT support HTML line breaks. Instead: use a semicolon (;) or the word "and" to separate multiple values in a single cell. If a cell has a list, pick the most important value or abbreviate.
- Naming conventions must be applied throughout — module names, asset names, queue names, argument names must follow the conventions defined in Sheet 2.
- Pre-populate standard entries: the four standard UiPath packages in Sheet 7, standard machine infrastructure defaults, standard compliance/security checklist items.
- Flag gaps clearly with [⚠ TO BE COMPLETED BY SA/DEV] for any row or field still pending.
- Never invent module names, asset values, queue schemas, or compliance answers — only use what the assessment data provides.
- The document should read as a technically credible, handover-ready artifact — precise enough for a developer to troubleshoot without reading the code.

DIAGRAM RULES:
- NEVER use Mermaid.js, ASCII art, or plain-text box diagrams anywhere in this document.
- For the current-state process flow in Sheet 3 (Diagram 1), output exactly this placeholder: {{BPMN_DIAGRAM}}
- For ALL other diagrams (proposed-state flow, architecture, module-level flows, ER diagrams, etc.), output a placeholder in EXACTLY this format:
  [[FLOWCHART: <a precise, specific description of what this diagram should visualise — include module names, step sequence, decision branches, system names, and flow direction>]]
- The placeholder will be replaced by a real AI-generated image at render time. Write the description as if briefing a visual designer.
- Example of a valid flowchart placeholder:
  [[FLOWCHART: High-level to-be automation flow. Left to right. Trigger: Orchestrator time-based trigger fires. Dispatcher reads SAP queue. Performer loop: Login SAP → Download invoice → OCR extract → 3-way match → Decision: Pass? Yes → Post to SAP FB60. No → Business exception queue. End: Daily Orchestrator summary report sent to manager.]]

Generate the COMPLETE SDD using the following EXACT 10-sheet structure:

---

# [Process Name] — Solution Design Document (SDD)
## RPA Automation Project | UiPath

---

## Sheet 1: Introduction

This Solution Design Document (SDD) describes the technical design of the UiPath RPA automation for the **[Process Name]** process within **[Business Unit / Org]**.

**Purpose:** This document serves as the primary communication artifact between the RPA Developer, RPA COE, Orchestrator/Bot Controller, and support teams. It provides sufficient technical detail for any team member to monitor, troubleshoot, and enhance the automation without needing to read the source code.

**Scope:** The SDD covers the robustness, scalability, efficiency, replicability, and reusability of all automation components for this process.

**Authors:** RPA Solution Architect + RPA Developer (see Sheet 2 for key contacts).
**Reviewer:** Solution Architect reviews prior to operations handover.
**Document Naming Convention:** [BusinessProcessName]_SDD_v[Version]

**Intended Audience:**
- RPA Centre of Excellence (COE)
- IT Support / RPA Supervisor
- RPA Developers (for troubleshooting and change requests)
- Business Analysts

> This is a **living document**. It must be updated whenever the automation design changes. All changes must reference the originating PDD or change request.

---

## Sheet 2: Process Overview

### Basic Info

| Field | Details |
|-------|---------|
| Process Name | [from assessment data] |
| Business Unit | [from assessment data] |
| Brief Process Summary | [from executive_summary in assessment data] |

### Architectural Structure

| Field | Details |
|-------|---------|
| Master Project Name | [⚠ TO BE COMPLETED BY SA/DEV] |
| Robot / Automation Type | Unattended (derive from assessment data; flag if unclear) |
| Robot(s) Used | [⚠ TO BE COMPLETED BY SA/DEV — names earmarked for this process] |
| Process Name(s) | [⚠ TO BE COMPLETED BY SA/DEV — Dispatcher, Performer, or single process] |
| Orchestrator Used? | Yes — for scheduling, queue management, asset storage, and monitoring |
| Scalable (parallel bots)? | [⚠ TO BE COMPLETED BY SA/DEV] |
| Process Type | Windows |
| UiPath Version | [⚠ TO BE COMPLETED BY SA/DEV] |

### Pre-Requisites

| System | Access Level | Description | Role |
|--------|-------------|-------------|------|
[Populate from systems_mentioned in the assessment. For each system, specify access level (Read/Write/Admin) and the bot role. Use n/a for unknown fields. Add rows as needed.]

### Application Installations

| Application | Version | Notes | Type (Thick/Thin) | Authentication | Access Method | App Location |
|------------|---------|-------|-------------------|---------------|--------------|-------------|
[Populate from systems_mentioned. Authentication: SSO / SAML / Basic Auth. Access Method: API / UI. Use n/a for unknown fields.]

### Machine Infrastructure

| Hardware | Setup | Version |
|----------|-------|---------|
| CPU | Minimum: Intel Core i5 or equivalent | [⚠ TO BE COMPLETED BY SA/DEV] |
| RAM | Minimum: 8 GB | [⚠ TO BE COMPLETED BY SA/DEV] |
| Windows OS | Windows 10 / Windows Server 2019 or later | [⚠ TO BE COMPLETED BY SA/DEV] |
| Storage | Minimum: 50 GB available | [⚠ TO BE COMPLETED BY SA/DEV] |

### Naming Conventions

| # | Convention | Example |
|---|-----------|---------|
| 1 | Variables | camelCase (e.g. invoiceNumber, vendorName) |
| 2 | Arguments | direction_ArgumentName (e.g. in_FileExists, io_Config, out_InvoiceData) |
| 3 | Tenant | [ProcessCode]_Tenant |
| 4 | Queue | [ProcessCode]_[QueuePurpose] (e.g. COE134_InvoiceQueue) |
| 5 | Folder | Department/AutomationMain/SubProcess |
| 6 | Modules | Action/App_FunctionalityName (e.g. ACME_Login, Filter_OldDataFromSharepoint) |
| 7 | Process | [ProcessCode]_[ProcessName] |
| 8 | Package | [Org].[ProcessName].[Component] |
| 9 | Assets | [ProcessCode]_AssetName (e.g. COE134_EpicLoginCredentials) |

### Project Organization

```
[ProcessName]/
├── ReusableModules/          # Shared/reusable workflows
├── [AppName]Layer/           # Per-application UI automation workflows
├── BusinessLogic/            # Core process logic workflows
├── Framework/                # REFramework or custom framework files
└── Data/
    ├── Config/               # Config files and environment variables
    ├── Templates/            # Input/output file templates
    └── ExceptionScreenshots/ # Auto-captured on error
```

---

## Sheet 3: Diagrams

### Diagram 1: Current State Process Flow (As-Is)

{{BPMN_DIAGRAM}}

**Written Description:**
[Provide a step-by-step numbered narrative of the current as-is process flow derived from the assessment data's process_map and scored_steps. Include trigger, each manual step, system touched, and output.]

### Diagram 2: High-Level Process Flow Design (To-Be)

[Use Mermaid flowchart LR to show the proposed automated flow. Include: Trigger → Dispatcher (if used) → Queue → Performer → Success/Exception branches → Notification/Reporting. Derive from the opportunities and priority_targets in the assessment data.]

**Step-by-step numbered process flow:**
[Supplement the diagram with a numbered text description: trigger conditions, retry logic, alert/notification mechanism, framework fit (REFramework / DU Framework / custom), rerun triggers, and scalability settings.]

**Notes:**
- Retry logic: [describe from assessment data or flag [⚠ TO BE COMPLETED BY SA/DEV]]
- Alert mechanism: [email / Orchestrator alert / custom notification]
- Scalability: [can multiple performers run in parallel? flag if unknown]
- Trigger disable for debugging: [⚠ TO BE COMPLETED BY SA/DEV]

### Diagram 3: Module-Level Flow — Login Example

[Use Mermaid flowchart TD to show a representative module-level flow, e.g., Login to the primary application. Include: Open browser/app → Enter credentials → Verify login → Handle failure → Return output argument. Adapt to the actual primary application from systems_mentioned.]

**Design Change Log:**

| Date | Change Description | PDD / CR Reference | Changed By |
|------|-------------------|-------------------|-----------|
| [today's date] | Initial SDD created from AP assessment | AI-Generated | AP Discovery Agent |

---

## Sheet 4: Design (Workflow Breakdown)

| Module/Workflow Name | Invoked In | Pre-Condition | Post-Condition | Description | Is Reusable? | Arguments (in/out/in-out) | Additional Notes | Exceptions | PDD Ref | Estimate | Developer | Status | Tested |
|---------------------|-----------|--------------|---------------|-------------|-------------|--------------------------|-----------------|-----------|---------|---------|----------|--------|--------|
[Populate one row per module. Derive module names from the process_map and scored_steps in assessment data. Follow naming convention: Action/App_FunctionalityName. Arguments follow convention: in_Name / out_Name / io_Name. Status = Not Started by default. Flag rows needing detail with [⚠ TO BE COMPLETED BY SA/DEV]. Add all key modules: initialization, per-app login, core business logic steps, exception handling, reporting/cleanup.]

---

## Sheet 5: Asset Management

| Asset Name | Description | Folder Path | Asset Type | Default Value | UAT Value | PROD Value | Additional Notes |
|-----------|-------------|------------|-----------|--------------|----------|-----------|-----------------|
[Populate from systems_mentioned (credentials per system) and assessment context. Follow naming convention: [ProcessCode]_AssetName. Asset Type: String / Bool / Int / Credential. Notes column must explain how each asset controls specific functionality in the code. Flag unknown values with [⚠ TO BE COMPLETED BY SA/DEV].]

---

## Sheet 6: Queues & Triggers

### Triggers

| Process | Type | Recurrence | Folder Path | Calendar / Notes |
|---------|------|-----------|------------|-----------------|
| [ProcessName]_Dispatcher | Time-based | [⚠ TO BE COMPLETED BY SA/DEV] | [⚠ Folder Path] | [⚠ Notes on calendar, holidays, peak periods] |
| [ProcessName]_Performer | Queue Trigger | When queue has items | [⚠ Folder Path] | [⚠ Min/Max bots, scale settings] |

### Queues

| Queue Name | Folder Path | Details | Queue Item Schema |
|-----------|------------|---------|-----------------|
[Populate based on the process flow. Queue name follows convention: [ProcessCode]_[Purpose]. Details must include: retry count, priority, unique reference field, purpose. Queue Item Schema must list all fields in the queue item (JSON structure or field table). If multiple item types exist, list each separately.]

### Orchestrator Folder Structure

| Main Folder | Parent Folder | Child Folder |
|------------|--------------|-------------|
| [Department] | AutomationMain | [ProcessName] |
| [Department] | AutomationMain | [ProcessName]/Shared |

---

## Sheet 7: Code Dependencies

### Packages

| SL# | Package Name | Version | External Library? | Is Approved? |
|-----|-------------|---------|-------------------|-------------|
| 1 | UiPath.Excel.Activities | [⚠ Latest stable] | No | Yes |
| 2 | UiPath.System.Activities | [⚠ Latest stable] | No | Yes |
| 3 | UiPath.Testing.Activities | [⚠ Latest stable] | No | Yes |
| 4 | UiPath.UIAutomation.Activities | [⚠ Latest stable] | No | Yes |
[Add any additional packages required based on systems_mentioned (e.g. UiPath.Web.Activities for web apps, UiPath.PDF.Activities for PDF extraction). Mark external libraries with Yes and add approval reference.]

### Other Dependencies / Appendix

| SL# | Other Dependencies / Appendix | Object | Additional Notes |
|-----|------------------------------|--------|-----------------|
| 1 | Assessment Data | JSON | AI-generated AP process analysis — source data for this SDD |
[Add any API definitions, Swagger files, embedded objects, or supporting documents referenced in the assessment data.]

---

## Sheet 8: Compliance & Security

### Compliance Considerations

| Compliance Item | Applicable? | Details |
|----------------|------------|---------|
| Sarbanes Oxley (SOX) | [Yes/No — derive from assessment context] | [If Yes: describe controls; if No: n/a] |
| HIPAA | [Yes/No] | [If Yes: describe PHI handling; if No: n/a] |
| FERC Standards of Conduct | [Yes/No] | n/a unless energy/utility process |
| PII (Principle of Least Privilege) | [Yes — derive from systems accessing personal data] | [Describe PoLP implementation: bot accesses only required fields] |
| Other | [⚠ TO BE COMPLETED BY SA/DEV] | [Any regional/industry-specific compliance] |
| File Sharing | [Yes/No] | [Describe how files are shared between bot and systems] |

### Security Checklist

| Security Item | Status | Notes |
|--------------|--------|-------|
| Logging any sensitive data? | [⚠ TO BE COMPLETED BY SA/DEV] | No PII should appear in plain-text logs |
| Loading PII to Orchestrator / Storage Bucket? | [⚠ TO BE COMPLETED BY SA/DEV] | If yes, must be encrypted |
| Credentials stored in approved vaults? | Yes | All credentials stored as Orchestrator Credential Assets |
| Sensitive information stored in clear text? | No | All sensitive data encrypted at rest |
| Extracting / reading fields that aren't necessary? | [⚠ TO BE COMPLETED BY SA/DEV] | Apply PoLP — read only required fields |

---

## Sheet 9: Reporting & Dashboarding

### Operational KPIs

| KPI | Description | Target / Benchmark |
|-----|-------------|-------------------|
| Robot Utilization | % of time robot is actively processing vs idle | >80% during scheduled window |
| Success Rate | % of transactions processed without exception | >95% |
| Volume / Transactions | Total items processed per run / per day | [derive from assessment data volumes] |
| Average Duration per Transaction | Time taken per queue item | [derive from assessment AHT data] |
| Error / Exception Rate | % of items hitting business or system exceptions | <5% |

### Business KPIs

| KPI | Description | Target / Benchmark |
|-----|-------------|-------------------|
| ROI | Cost saved vs automation cost | [derive from assessment ROI data if available] |
| FTE Hours Saved | Hours freed per week/month | [derive from assessment FTE × AHT data] |
| Process Cycle Time Reduction | Reduction in end-to-end processing time | [derive from assessment data] |
| Financial Savings | Cost per invoice manual vs automated | [derive from assessment hourly_rate data if available] |

### Industry / Process-Specific KPIs

| KPI | Description | Target |
|-----|-------------|--------|
| Invoice Straight-Through Rate | % of invoices processed without human touch | >90% |
| 3-Way Match Success Rate | % of POs matched to invoices and receipts automatically | >85% |
| Exception Resolution Time | Avg time to resolve a flagged exception | <24 hours |
[Add or remove KPIs based on the specific process domain from the assessment data.]

### Value Reporting Summary

Operational + industry-specific KPIs will be tracked via UiPath Orchestrator dashboards and supplemented by a weekly email summary report to the Process Owner and RPA COE.

---

## Sheet 10: UAT Configuration

### Test Cases

| ID | Category | Test Case | Inputs | Expected Outputs | Additional Pass Condition | Status | Tests Count | Tests Success | Fail Reasons | Business Owner |
|----|---------|-----------|--------|-----------------|--------------------------|--------|------------|--------------|-------------|---------------|
| TC-001 | End-To-End | Full happy-path processing of a standard invoice | Valid invoice, matching PO and GR | Invoice processed successfully, payment queued | No exceptions raised; audit log entry created | Not Run | [⚠] | [⚠] | — | [⚠ TO BE DEFINED BY QA/BA] |
| TC-002 | Business Exception | Invoice with no matching PO | Invoice with non-existent PO number | Transaction flagged as business exception | Bot notifies SME; queue item marked as Failed | Not Run | [⚠] | [⚠] | — | [⚠ TO BE DEFINED BY QA/BA] |
| TC-003 | Business Exception | Duplicate invoice detected | Invoice number already in system | Duplicate flagged; no double payment | Alert sent; original invoice untouched | Not Run | [⚠] | [⚠] | — | [⚠ TO BE DEFINED BY QA/BA] |
| TC-004 | Process Specific | Invoice amount exceeds approval threshold | Invoice > threshold value | Routed to manual approval workflow | Bot pauses processing; handoff documented | Not Run | [⚠] | [⚠] | — | [⚠ TO BE DEFINED BY QA/BA] |
| TC-005 | Business Exception | 3-way match failure | Mismatched quantities/amounts | Exception raised; item sent to exceptions queue | Discrepancy details logged | Not Run | [⚠] | [⚠] | — | [⚠ TO BE DEFINED BY QA/BA] |
| TC-006 | Edge Case | System unavailable at process start | Target app offline | Bot retries per retry policy, then alerts | Max retries exhausted → graceful shutdown | Not Run | [⚠] | [⚠] | — | [⚠ TO BE DEFINED BY QA/BA] |
| TC-007 | Edge Case | Empty queue at trigger time | No items in Orchestrator queue | Process exits gracefully with log entry | No error raised; scheduled run completes | Not Run | [⚠] | [⚠] | — | [⚠ TO BE DEFINED BY QA/BA] |
[Add additional test cases derived from the pain_points and exceptions identified in the assessment data. All Status values default to "Not Run".]

---

Here is the assessment data to base this SDD on:
{context}
"""

# ── Flowchart Image Injection ───────────────────────────────────────────────────

def inject_flowcharts(content: str) -> str:
    """Replace [[FLOWCHART: ...]] placeholders with AI-generated SVG diagrams
    using a 3-stage LangChain-style pipeline.

    PIPELINE
    --------
    Stage 1 - Process Extractor:  reads the FULL document and extracts a
              structured process model (steps, decisions, exceptions, actors,
              edges) from the section the placeholder refers to.
    Stage 2 - Layout Planner:     converts that structured model into a
              precise SVG layout (viewBox, node positions, sizing).
    Stage 3 - SVG Writer:         turns the layout into polished SVG XML.

    This produces diagrams that capture EVERY exception, gateway and
    validation mentioned anywhere in the PDD/SDD - not just the happy path.

    Each generated SVG is run through `_clean_and_wrap_svg`:
      - Strips stray fences and <?xml ...?>.
      - Replaces invalid height/width="auto" with "100%".
      - Collapses ALL whitespace to a single space (so markdown does not split
        the SVG into separate <pre><code> fragments at blank lines).
      - Wraps in <figure class="doc-flowchart">.

    Also defensively rewraps raw <svg> blocks the AI may emit outside any
    [[FLOWCHART:]] placeholder.
    """
    import re
    from utils.gemini_client import generate_flowchart_svg, extract_process_for_diagram

    def _clean_and_wrap_svg(svg_text: str) -> str:
        s = svg_text.strip()
        s = re.sub(r'^<\?xml[^?]*\?>\s*', '', s)
        s = re.sub(r'^```(?:svg|xml|html)?\s*', '', s, flags=re.IGNORECASE)
        s = re.sub(r'\s*```\s*$', '', s)
        s = re.sub(r'\bheight\s*=\s*["\']auto["\']', 'height="100%"', s, flags=re.IGNORECASE)
        s = re.sub(r'\bwidth\s*=\s*["\']auto["\']', 'width="100%"', s, flags=re.IGNORECASE)
        s_oneline = re.sub(r'\s+', ' ', s).strip()
        return (
            f'\n\n<figure class="doc-flowchart" style="margin:2rem 0;text-align:center;">'
            f'{s_oneline}'
            f'</figure>\n\n'
        )

    pattern = re.compile(r'\[\[FLOWCHART:\s*(.+?)\]\]', re.DOTALL)
    extraction_cache: dict = {}

    def replace_placeholder(match: re.Match) -> str:
        description = match.group(1).strip()
        print(f"[inject_flowcharts] Building diagram via 3-stage pipeline for: {description[:80]}...")
        cache_key = description.lower()[:120]
        if cache_key in extraction_cache:
            structure = extraction_cache[cache_key]
        else:
            try:
                structure = extract_process_for_diagram(content, description)
            except Exception as e:
                print(f"[inject_flowcharts] Stage 1 (extractor) errored: {e}")
                structure = None
            extraction_cache[cache_key] = structure
        try:
            svg_content = generate_flowchart_svg(description, process_structure=structure)
        except Exception as e:
            print(f"[inject_flowcharts] generate_flowchart_svg errored: {e}")
            svg_content = None
        if svg_content:
            return _clean_and_wrap_svg(svg_content)
        return (
            f"\n> **Diagram:** *{description}*  \n"
            f"> *(AI-generated flowchart image could not be rendered - "
            f"please add the diagram manually.)*\n"
        )

    content = pattern.sub(replace_placeholder, content)

    fence_re = re.compile(
        r'```(?:svg|xml|html)\s*\n(<svg\b[\s\S]*?</svg>)\s*\n```',
        re.IGNORECASE,
    )
    content = fence_re.sub(lambda m: _clean_and_wrap_svg(m.group(1)), content)

    # Step 3: Defensive - re-wrap RAW <svg>...</svg> blocks that aren't
    # already inside a <figure>. Python's re module rejects variable-width
    # look-behinds, so we use finditer + a small prefix check instead.
    raw_svg_re = re.compile(r'<svg\b[\s\S]*?</svg>', re.IGNORECASE)
    parts: list = []
    last_end = 0
    for m in raw_svg_re.finditer(content):
        parts.append(content[last_end:m.start()])
        svg_block = m.group(0)
        # Look at the immediate prefix - if there's an unclosed <figure ...>
        # right before this svg, it's already wrapped, leave it alone.
        prefix = content[max(0, m.start() - 200):m.start()]
        last_open = prefix.rfind('<figure')
        last_close = prefix.rfind('</figure>')
        already_in_figure = last_open > last_close
        if '\n' not in svg_block or already_in_figure:
            parts.append(svg_block)
        else:
            parts.append(_clean_and_wrap_svg(svg_block).strip())
        last_end = m.end()
    parts.append(content[last_end:])
    content = ''.join(parts)

    return content

class DocumentRequest(BaseModel):
    doc_type: str = "pdd"  # "pdd" or "sdd"
    force: bool = False

class DocumentUpdateRequest(BaseModel):
    doc_type: str = "pdd"
    instruction: str
    selected_text: Optional[str] = None

class DocumentManualUpdateRequest(BaseModel):
    content: str

@app.post("/api/history/{assessment_id}/upload-template")
async def upload_template(
    assessment_id: str,
    doc_type: str = Form("pdd"),
    file: UploadFile = File(...),
    user_id: str = Depends(get_current_user),
):
    """Upload a custom PDD/SDD Word template (.docx) for an assessment.

    Once uploaded, the generated document follows this template's section
    structure, and the downloaded .docx reuses the template's design (logo,
    header/footer, fonts and styles).
    """
    doc_type = doc_type.lower()
    if doc_type not in ("pdd", "sdd"):
        raise HTTPException(status_code=400, detail="doc_type must be 'pdd' or 'sdd'")
    if not file.filename or not file.filename.lower().endswith(".docx"):
        raise HTTPException(status_code=400, detail="Template must be a .docx file")

    assessment = get_assessment(assessment_id, user_id)
    if not assessment:
        raise HTTPException(status_code=404, detail="Assessment not found")

    data = await file.read()
    path = save_template(assessment_id, doc_type, data)
    try:
        outline = build_template_outline(path)
        sections = [ln.strip() for ln in outline.split("\n") if ln.strip()]
    except Exception as e:
        print(f"[upload_template] outline parse failed: {e}")
        sections = []
    return {
        "status": "ok",
        "doc_type": doc_type,
        "filename": file.filename,
        "section_count": len(sections),
        "sections": sections,
    }


@app.get("/api/history/{assessment_id}/template-status")
async def template_status(assessment_id: str, user_id: str = Depends(get_current_user)):
    """Report whether a custom template has been uploaded for each doc type."""
    assessment = get_assessment(assessment_id, user_id)
    if not assessment:
        raise HTTPException(status_code=404, detail="Assessment not found")
    return {
        "pdd": has_template(assessment_id, "pdd"),
        "sdd": has_template(assessment_id, "sdd"),
    }


@app.delete("/api/history/{assessment_id}/template")
async def remove_template(assessment_id: str, doc_type: str = "pdd", user_id: str = Depends(get_current_user)):
    """Remove a previously uploaded template so generation reverts to default."""
    assessment = get_assessment(assessment_id, user_id)
    if not assessment:
        raise HTTPException(status_code=404, detail="Assessment not found")
    deleted = delete_template(assessment_id, doc_type.lower())
    return {"status": "ok", "deleted": deleted, "doc_type": doc_type.lower()}


@app.post("/api/history/{assessment_id}/generate-document")
async def generate_document(assessment_id: str, payload: DocumentRequest, user_id: str = Depends(get_current_user)):
    """Generate a PDD or SDD document for an assessment."""
    try:
        doc_type = payload.doc_type.lower()
        if doc_type not in ("pdd", "sdd"):
            raise HTTPException(status_code=400, detail="doc_type must be 'pdd' or 'sdd'")

        assessment = get_assessment(assessment_id, user_id)
        if not assessment:
            raise HTTPException(status_code=404, detail="Assessment not found")

        # Check if document already exists in MongoDB
        if not payload.force:
            documents = assessment.get("documents", {})
            existing = documents.get(doc_type)
            if existing:
                return {"doc_type": doc_type, "content": existing, "cached": True}

        # Build context from assessment data
        # Exclude fields that are either:
        #   (a) Too large to safely fit in a prompt (original_transcript, markdown_report, bpmn_xml, chat_history, audio_scripts)
        #   (b) Internal metadata not useful for doc generation (_id, user_id)
        #   (c) Already applied post-generation (documents)
        HEAVY_FIELDS = {
            'bpmn_xml', 'chat_history', '_id', 'user_id', 'audio_scripts', 'documents',
            # These two are the main token bombs — original_transcript can be 500k+ chars
            # from a long audio/video file, and markdown_report is the full analysis report.
            # The PDD/SDD templates only need the structured fields below.
            'original_transcript', 'markdown_report',
        }
        context_data = {k: v for k, v in assessment.items() if k not in HEAVY_FIELDS}
        context_str = json.dumps(context_data, indent=2)

        # Safety guard: if remaining context is still very large (e.g. many scored_steps),
        # truncate to ~200,000 chars (~50,000 tokens) which is well within the 1M limit
        # when combined with the template (~6,000 tokens).
        MAX_CONTEXT_CHARS = 200_000
        if len(context_str) > MAX_CONTEXT_CHARS:
            print(f"[generate_document] context_str too large ({len(context_str)} chars) — truncating to {MAX_CONTEXT_CHARS}")
            context_str = context_str[:MAX_CONTEXT_CHARS] + "\n... [truncated for length] ..."

        # Choose the right template / prompt
        if has_template(assessment_id, doc_type):
            # User uploaded their own template -> generate content that follows
            # its exact section structure. The visual design (logo, header/
            # footer, fonts) is applied at download time by rendering this
            # content into a copy of the uploaded .docx.
            outline = build_template_outline(get_template_path(assessment_id, doc_type))
            prompt = (
                "You are a senior RPA documentation specialist. Generate a COMPLETE "
                + doc_type.upper()
                + " in GitHub-flavoured Markdown that follows EXACTLY the client's "
                "official template section structure shown below.\n\n"
                "RULES:\n"
                "- Reproduce the template's headings EXACTLY (same wording, numbering and order).\n"
                "- Use '#' for top-level sections, '##' for sub-sections and '###' for deeper levels, matching the outline indentation.\n"
                "- Fill every section with specific content from the assessment data below. Where data is genuinely missing, write '[TO BE COMPLETED BY SME]'.\n"
                "- Use Markdown tables for tabular data. Do not use <br> inside tables; separate items with ';'.\n"
                "- For the main end-to-end process flow, output a single line containing exactly [[BPMN_DIAGRAM]].\n"
                "- For any other diagram, output a single line [[FLOWCHART: short description]].\n\n"
                "TEMPLATE SECTION STRUCTURE (reproduce these headings exactly, in order):\n"
                + outline
                + "\n\nASSESSMENT DATA (JSON):\n"
                + context_str
            )
        else:
            template = PDD_TEMPLATE if doc_type == "pdd" else SDD_TEMPLATE
            prompt = template.format(context=context_str)

        # Generate with Gemini 3.5 Flash
        content = call_gemini_text(
            prompt=prompt,
            system_prompt="You are a senior document specialist. Generate detailed, professional documents with proper markdown formatting. Be specific — use the actual data provided, not generic placeholders.",
            model="gemini-3.5-flash"
        )

        # ── Auto-generate BPMN if missing ────────────────────────────────
        bpmn_xml = assessment.get("bpmn_xml", "")
        if not bpmn_xml:
            scored_steps = assessment.get("scored_steps", [])
            if scored_steps:
                print(f"[generate_document] No BPMN diagram found — auto-generating from {len(scored_steps)} scored steps...")
                try:
                    bpmn_xml = _generate_bpmn_xml(scored_steps)
                    if bpmn_xml:
                        update_assessment_bpmn(assessment_id, user_id, bpmn_xml)
                        assessment["bpmn_xml"] = bpmn_xml
                        print(f"[generate_document] BPMN auto-generated and saved ({len(bpmn_xml)} chars)")
                except Exception as e:
                    print(f"[generate_document] BPMN auto-generation failed (non-fatal): {e}")

        # Step 1: Inject pre-generated BPMN XML into the placeholder
        bpmn_xml = assessment.get("bpmn_xml", "")
        if bpmn_xml:
            bpmn_markdown = f"```bpmn\n{bpmn_xml}\n```"
            # Python's .format() escapes {{ }} into { }, so the model often outputs {BPMN_DIAGRAM}
            if "{{BPMN_DIAGRAM}}" in content:
                content = content.replace("{{BPMN_DIAGRAM}}", bpmn_markdown)
            elif "{BPMN_DIAGRAM}" in content:
                content = content.replace("{BPMN_DIAGRAM}", bpmn_markdown)
            elif "[[BPMN_DIAGRAM]]" in content:
                content = content.replace("[[BPMN_DIAGRAM]]", bpmn_markdown)

        # Step 2: Replace [[FLOWCHART: ...]] placeholders with Gemini-generated images
        content = inject_flowcharts(content)

        # Save to MongoDB
        update_assessment_document(assessment_id, user_id, doc_type, content)

        return {"doc_type": doc_type, "content": content, "cached": False}
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


@app.post("/api/history/{assessment_id}/update-document")
async def update_document(assessment_id: str, payload: DocumentUpdateRequest, user_id: str = Depends(get_current_user)):
    """Update a PDD/SDD document based on user instruction."""
    try:
        doc_type = payload.doc_type.lower()
        if doc_type not in ("pdd", "sdd"):
            raise HTTPException(status_code=400, detail="doc_type must be 'pdd' or 'sdd'")

        assessment = get_assessment(assessment_id, user_id)
        if not assessment:
            raise HTTPException(status_code=404, detail="Assessment not found")

        documents = assessment.get("documents", {})
        current_doc = documents.get(doc_type)
        if not current_doc:
            raise HTTPException(status_code=400, detail=f"No {doc_type.upper()} exists yet. Generate one first.")

        doc_label = "Process Design Document (PDD)" if doc_type == "pdd" else "Solution Design Document (SDD)"

        import re

        # ── Strip heavy embedded content before sending to Gemini ──────────────
        # The document stores flowchart diagrams as raw <figure> SVG HTML blocks.
        # Each one can be 5,000–15,000 chars — enough to blow the 1M token limit.
        # Strategy: replace each diagram block with a tiny [[EMBEDDED_IMAGE_N]]
        # placeholder, save the originals, then restore them after editing.
        # (Any *new* [[FLOWCHART: ...]] written by Gemini will be regenerated.)

        embedded_images: list[str] = []  # ordered list of original content

        def _stash_image(m: re.Match) -> str:
            idx = len(embedded_images)
            embedded_images.append(m.group(0))
            return f"[[EMBEDDED_IMAGE_{idx}]]"  # tiny placeholder (~20 tokens)

        # Safe-keep BPMN diagram (large XML block)
        bpmn_xml = assessment.get("bpmn_xml", "")
        bpmn_block = ""
        if bpmn_xml:
            bpmn_block = f"```bpmn\n{bpmn_xml}\n```"
            # Replace bpmn block with placeholder first
            _bpmn_re = r"```bpmn\s+[\s\S]*?```"
            current_doc_for_gemini = re.sub(_bpmn_re, "{{BPMN_DIAGRAM}}", current_doc)
        else:
            current_doc_for_gemini = current_doc

        # Stash raw <figure class="doc-flowchart"> SVG HTML blocks (new format)
        _figure_re = re.compile(
            r'<figure[^>]*class=["\']doc-flowchart["\'][^>]*>[\s\S]*?</figure>',
            re.MULTILINE
        )
        current_doc_for_gemini = _figure_re.sub(_stash_image, current_doc_for_gemini)

        # Also stash any legacy markdown images with data: URLs (old format safety net)
        _data_url_re = re.compile(
            r'!\[([^\]]*)\]\(data:image/[^;]+;base64,[A-Za-z0-9+/=\r\n]+\)',
            re.MULTILINE
        )
        current_doc_for_gemini = _data_url_re.sub(_stash_image, current_doc_for_gemini)

        if embedded_images:
            print(f"[update_document] Stripped {len(embedded_images)} diagram block(s) to save tokens.")

        selected_text_part = ""
        if payload.selected_text and payload.selected_text.strip():
            selected_text_part = f"\nSpecifically, in the following section/text of the document:\n\"{payload.selected_text}\"\n"

        prompt = f"""You are a senior document specialist. Below is the current {doc_label}.

The user wants you to make the following change:
"{payload.instruction}"
{selected_text_part}
Apply the requested change to the document. Return the COMPLETE updated document with ALL sections intact — do not omit or summarize any sections that were not changed. Only modify the sections that the user's instruction/selection refers to. Keep the same professional markdown formatting.
IMPORTANT: Do NOT modify these placeholders — leave them exactly as they appear:
- '{{BPMN_DIAGRAM}}'
- Any '[[EMBEDDED_IMAGE_N]]' placeholders (where N is a number)

Current document:
{current_doc_for_gemini}
"""

        updated_content = call_gemini_text(
            prompt=prompt,
            system_prompt="You are a senior document specialist. Apply the requested changes carefully. Return the complete updated document — never truncate or summarize unchanged sections.",
            model="gemini-3.5-flash"
        )

        # ── Restore stashed images ─────────────────────────────────────────────
        for idx, original_img in enumerate(embedded_images):
            updated_content = updated_content.replace(f"[[EMBEDDED_IMAGE_{idx}]]", original_img)

        # ── Re-inject BPMN block ───────────────────────────────────────────────
        if bpmn_xml and bpmn_block:
            if "{{BPMN_DIAGRAM}}" in updated_content:
                updated_content = updated_content.replace("{{BPMN_DIAGRAM}}", bpmn_block)
            elif "{BPMN_DIAGRAM}" in updated_content:
                updated_content = updated_content.replace("{BPMN_DIAGRAM}", bpmn_block)
            elif "[[BPMN_DIAGRAM]]" in updated_content:
                updated_content = updated_content.replace("[[BPMN_DIAGRAM]]", bpmn_block)

        # ── Generate SVG images for any NEW [[FLOWCHART: ...]] placeholders ────
        updated_content = inject_flowcharts(updated_content)

        # Replace the old document in MongoDB
        update_assessment_document(assessment_id, user_id, doc_type, updated_content)

        return {"doc_type": doc_type, "content": updated_content}
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


@app.put("/api/history/{assessment_id}/document/{doc_type}")
async def manual_update_document(assessment_id: str, doc_type: str, payload: DocumentManualUpdateRequest, user_id: str = Depends(get_current_user)):
    """Manually update the PDD or SDD document content."""
    try:
        doc_type = doc_type.lower()
        if doc_type not in ("pdd", "sdd"):
            raise HTTPException(status_code=400, detail="doc_type must be 'pdd' or 'sdd'")

        assessment = get_assessment(assessment_id, user_id)
        if not assessment:
            raise HTTPException(status_code=404, detail="Assessment not found")

        update_assessment_document(assessment_id, user_id, doc_type, payload.content)
        return {"doc_type": doc_type, "content": payload.content}
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


@app.post("/api/history/{assessment_id}/download-document")
async def download_document_docx(
    assessment_id: str,
    doc_type: str = "pdd",
    payload: dict = Body(None),
    user_id: str = Depends(get_current_user),
):
    """Download the PDD/SDD as a styled Word (.docx) file.

    Frontend POSTs `{ images: [{ type: 'bpmn'|'svg', png_base64: '...' }] }` -
    PNG snapshots of every rendered diagram captured from the DOM, in document
    order. We embed those PNGs in place of `bpmn` code blocks and
    `<figure class="doc-flowchart">` blocks so the Word doc has real images
    instead of raw XML/HTML markup.
    """
    try:
        assessment = get_assessment(assessment_id, user_id)
        if not assessment:
            raise HTTPException(status_code=404, detail="Assessment not found")

        documents = assessment.get("documents", {})
        content = documents.get(doc_type)
        if not content:
            raise HTTPException(status_code=400, detail=f"No {doc_type.upper()} document exists yet.")

        # If the user uploaded their own template for this doc type, render the
        # generated content INTO a copy of that template so the output keeps the
        # template logo, header/footer, fonts and section styling.
        _tpl = get_template_path(assessment_id, doc_type)
        if _tpl:
            _imgs = (payload or {}).get("images", []) if isinstance(payload, dict) else []
            _docx_bytes = render_markdown_into_template(_tpl, content, _imgs)
            return Response(
                content=_docx_bytes,
                media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                headers={"Content-Disposition": f'attachment; filename="AP_{doc_type.upper()}_Document.docx"'},
            )

        import io
        import re
        import base64
        import requests
        from docx import Document
        from docx.shared import Pt, Inches, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        # Frontend-captured diagram PNGs, in document order.
        images = (payload or {}).get("images", []) if isinstance(payload, dict) else []
        bpmn_images = [im for im in images if im.get("type") == "bpmn" and im.get("png_base64")]
        svg_images = [im for im in images if im.get("type") == "svg" and im.get("png_base64")]
        bpmn_idx = 0
        svg_idx = 0

        doc = Document()

        # Set default font
        style = doc.styles['Normal']
        font = style.font
        font.name = 'Calibri'
        font.size = Pt(11)
        font.color.rgb = RGBColor(0x33, 0x33, 0x33)

        # ── Page setup: narrower margins give wide tables & diagrams more room ──
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn
        for _section in doc.sections:
            _section.left_margin = Inches(0.7)
            _section.right_margin = Inches(0.7)
            _section.top_margin = Inches(0.8)
            _section.bottom_margin = Inches(0.8)
        _section0 = doc.sections[0]
        content_width = _section0.page_width - _section0.left_margin - _section0.right_margin
        content_height = _section0.page_height - _section0.top_margin - _section0.bottom_margin
        max_img_w = min(content_width, Inches(6.8))
        max_img_h = int(content_height * 0.9)

        def _fit_table(tbl):
            """Force a Word table to fit the page width and wrap long cell text,
            so no column is ever pushed off the right edge of the page."""
            tbl.autofit = True
            tbl.allow_autofit = True
            tblPr = tbl._tbl.tblPr
            for _el in tblPr.findall(qn('w:tblW')):
                tblPr.remove(_el)
            _tblW = OxmlElement('w:tblW')
            _tblW.set(qn('w:w'), '5000')        # 5000 fiftieths of a percent = 100%
            _tblW.set(qn('w:type'), 'pct')
            tblPr.append(_tblW)
            for _el in tblPr.findall(qn('w:tblLayout')):
                tblPr.remove(_el)
            _layout = OxmlElement('w:tblLayout')
            _layout.set(qn('w:type'), 'autofit')
            tblPr.append(_layout)
            for _row in tbl.rows:
                for _cell in _row.cells:
                    for _para in _cell.paragraphs:
                        for _run in _para.runs:
                            _run.font.size = Pt(8.5)
            try:
                _trPr = tbl.rows[0]._tr.get_or_add_trPr()
                _hdr = OxmlElement('w:tblHeader')
                _hdr.set(qn('w:val'), 'true')
                _trPr.append(_hdr)
            except Exception:
                pass

        def _add_fitted_picture(image_stream):
            """Insert an image centered and scaled to fit within the page, so it is
            never wider than the margins nor taller than one page (which is what
            causes diagrams to be split across two pages)."""
            data = image_stream.getvalue() if hasattr(image_stream, 'getvalue') else image_stream.read()
            try:
                from PIL import Image as _PILImage
                with _PILImage.open(io.BytesIO(data)) as _im:
                    _pw, _ph = _im.size
            except Exception:
                _pw, _ph = (0, 0)
            _p = doc.add_paragraph()
            _p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            _run = _p.add_run()
            if _pw and _ph:
                _aspect = _ph / _pw
                if int(max_img_w * _aspect) <= max_img_h:
                    _run.add_picture(io.BytesIO(data), width=max_img_w)
                else:
                    _run.add_picture(io.BytesIO(data), height=max_img_h)
            else:
                _run.add_picture(io.BytesIO(data), width=max_img_w)

        # Parse markdown line by line and build the Word document
        # Bulletproof flatten of <figure class="doc-flowchart">...</figure>
        # blocks into a single sentinel line, BEFORE we tokenize by '\n'.
        # The Gemini-generated SVG inside almost always contains newlines,
        # which used to make the figure span many lines and let the SVG XML
        # leak through to the paragraph branch as raw text.
        _flowchart_re = re.compile(
            r'<figure[^>]*class=["\']doc-flowchart["\'][^>]*>[\s\S]*?</figure>',
            re.MULTILINE
        )
        content = _flowchart_re.sub('\n[[FLOWCHART_FIGURE]]\n', content)

        lines = content.split('\n')
        i = 0
        in_table = False
        table_rows = []
        in_code_block = False
        code_language = ""
        code_lines = []

        while i < len(lines):
            line = lines[i]

            # Code block toggle
            if line.strip().startswith('```'):
                if not in_code_block:
                    in_code_block = True
                    code_language = line.strip()[3:].strip().lower()
                    code_lines = []
                else:
                    in_code_block = False
                    code_str = "\n".join(code_lines)
                    
                    if code_language in ['mermaid', 'bpmn']:
                        embedded = False
                        # Prefer the PNG snapshot the frontend captured from
                        # the rendered DOM - matches what the user sees and
                        # has no network dependency.
                        if code_language == 'bpmn' and bpmn_idx < len(bpmn_images):
                            try:
                                png_b64 = bpmn_images[bpmn_idx].get("png_base64", "")
                                if "," in png_b64:
                                    png_b64 = png_b64.split(",", 1)[1]
                                png_bytes = base64.b64decode(png_b64)
                                _add_fitted_picture(io.BytesIO(png_bytes))
                                embedded = True
                            except Exception as e:
                                print(f"[download_docx] BPMN PNG embed failed, falling back to Kroki: {e}")
                            bpmn_idx += 1

                        if not embedded:
                            try:
                                # Fallback: Kroki API renders Mermaid/BPMN as PNG.
                                import zlib
                                compressed = zlib.compress(code_str.encode('utf-8'), 9)
                                encoded = base64.urlsafe_b64encode(compressed).decode('utf-8')
                                engine = 'mermaid' if code_language == 'mermaid' else 'bpmn'
                                url = f"https://kroki.io/{engine}/png/{encoded}"

                                resp = requests.get(url, timeout=15)
                                if resp.status_code == 200:
                                    image_stream = io.BytesIO(resp.content)
                                    _add_fitted_picture(image_stream)
                                else:
                                    doc.add_paragraph(f"[{code_language.upper()} diagram failed to render online]", style='Normal')
                            except Exception as e:
                                doc.add_paragraph(f"[{code_language.upper()} render error: {e}]", style='Normal')
                    else:
                        p = doc.add_paragraph(code_str)
                        p.style = doc.styles['Normal']
                        run = p.runs[0] if p.runs else p.add_run(code_str)
                        run.font.name = 'Consolas'
                        run.font.size = Pt(9)
                        run.font.color.rgb = RGBColor(0x60, 0x60, 0x60)
                i += 1
                continue

            if in_code_block:
                code_lines.append(line)
                i += 1
                continue

            # Table detection
            if '|' in line and line.strip().startswith('|'):
                cells = [c.strip() for c in line.strip().strip('|').split('|')]
                # Skip separator rows (e.g., |---|---|)
                if all(set(c.strip()).issubset({'-', ':', ' '}) for c in cells):
                    i += 1
                    continue
                table_rows.append(cells)
                i += 1
                continue
            else:
                # Flush table if we were in one
                if table_rows:
                    cols = max(len(r) for r in table_rows)
                    tbl = doc.add_table(rows=len(table_rows), cols=cols)
                    tbl.style = 'Light Grid Accent 1'
                    for ri, row in enumerate(table_rows):
                        for ci, cell in enumerate(row):
                            if ci < cols:
                                tbl.rows[ri].cells[ci].text = cell
                                # Bold headers (first row)
                                if ri == 0:
                                    for run in tbl.rows[ri].cells[ci].paragraphs[0].runs:
                                        run.bold = True
                    _fit_table(tbl)
                    table_rows = []

            stripped = line.strip()

            # Headings
            if stripped.startswith('# ') and not stripped.startswith('##'):
                p = doc.add_heading(stripped[2:], level=0)
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                
                # Insert Native Table of Contents
                from docx.oxml import parse_xml
                from docx.oxml.ns import nsdecls
                
                doc.add_heading('Table of Contents', level=1)
                p_toc = doc.add_paragraph()
                run = p_toc.add_run()
                
                fldChar1 = parse_xml(r'<w:fldChar %s w:fldCharType="begin"/>' % nsdecls('w'))
                instrText = parse_xml(r'<w:instrText %s xml:space="preserve">TOC \o "1-3" \h \z \u</w:instrText>' % nsdecls('w'))
                fldChar2 = parse_xml(r'<w:fldChar %s w:fldCharType="separate"/>' % nsdecls('w'))
                fldChar3 = parse_xml(r'<w:fldChar %s w:fldCharType="end"/>' % nsdecls('w'))
                
                run._r.append(fldChar1)
                run._r.append(instrText)
                run._r.append(fldChar2)
                run._r.append(fldChar3)
                
                doc.add_page_break()
                
                # Attempt to set updateFields to true so Word calculates numbers on open
                try:
                    doc.settings.element.append(parse_xml(r'<w:updateFields %s w:val="true"/>' % nsdecls('w')))
                except Exception:
                    pass
                    
            elif stripped.startswith('## '):
                doc.add_heading(stripped[3:], level=1)
            elif stripped.startswith('### '):
                doc.add_heading(stripped[4:], level=2)
            elif stripped.startswith('#### '):
                doc.add_heading(stripped[5:], level=3)
            elif stripped == '[[FLOWCHART_FIGURE]]':
                # Sentinel placed by the regex pre-pass for every
                # <figure class="doc-flowchart">...</figure> block in the
                # markdown. Embed the matching frontend-captured PNG.
                if svg_idx < len(svg_images):
                    try:
                        png_b64 = svg_images[svg_idx].get("png_base64", "")
                        if "," in png_b64:
                            png_b64 = png_b64.split(",", 1)[1]
                        png_bytes = base64.b64decode(png_b64)
                        _add_fitted_picture(io.BytesIO(png_bytes))
                    except Exception as e:
                        print(f"[download_docx] SVG PNG embed failed: {e}")
                        doc.add_paragraph("[SVG diagram could not be embedded]", style='Normal')
                    svg_idx += 1
                else:
                    print(f"[download_docx] No frontend SVG capture for flowchart #{svg_idx + 1}")
                    doc.add_paragraph("[SVG diagram not captured - re-open the document and try again]", style='Normal')
                i += 1
                continue
            elif stripped.startswith('---'):
                doc.add_paragraph('').add_run().add_break()
            elif stripped.startswith('- ') or stripped.startswith('* '):
                text = stripped[2:]
                # Handle bold in list items
                text = re.sub(r'\*\*(.*?)\*\*', r'\1', text)
                doc.add_paragraph(text, style='List Bullet')
            elif stripped:
                # Regular paragraph — handle inline bold/italic
                p = doc.add_paragraph()
                # Split by bold markers
                parts = re.split(r'(\*\*.*?\*\*)', stripped)
                for part in parts:
                    if part.startswith('**') and part.endswith('**'):
                        run = p.add_run(part[2:-2])
                        run.bold = True
                    else:
                        p.add_run(part)

            i += 1

        # Flush any remaining table
        if table_rows:
            cols = max(len(r) for r in table_rows)
            tbl = doc.add_table(rows=len(table_rows), cols=cols)
            tbl.style = 'Light Grid Accent 1'
            for ri, row in enumerate(table_rows):
                for ci, cell in enumerate(row):
                    if ci < cols:
                        tbl.rows[ri].cells[ci].text = cell
            _fit_table(tbl)

        # Save to bytes
        fp = io.BytesIO()
        doc.save(fp)
        docx_bytes = fp.getvalue()

        filename = f"AP_{doc_type.upper()}_Document.docx"
        return Response(
            content=docx_bytes,
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={"Content-Disposition": f"attachment; filename={filename}"}
        )
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

# ── BPMN Generation System Prompt (Industry-Level) ──────────────────────────
BPMN_SYSTEM_PROMPT = """You are a certified BPMN 2.0 process architect with 15+ years of experience modelling
enterprise Accounts Payable (AP) processes for SAP, Oracle, QuickBooks, and paper-based environments.
You specialise in producing valid, executable-ready BPMN 2.0 XML that renders correctly in bpmn.io,
Camunda Modeler, and any standards-compliant BPMN viewer.

Your output is ALWAYS raw BPMN 2.0 XML wrapped in a ```xml code block — nothing else.
You NEVER produce prose, ASCII art, or Mermaid diagrams in this role.

═══════════════════════════════════════════════════════════════════
PART A — BPMN 2.0 ELEMENT RULES (MUST FOLLOW EXACTLY)
═══════════════════════════════════════════════════════════════════

A1. ROOT NAMESPACE DECLARATION (MANDATORY)
─────────────────────────────────────────
Every XML file MUST begin with the following exact namespace block:

<bpmn:definitions
    xmlns:bpmn="http://www.omg.org/spec/BPMN/20100524/MODEL"
    xmlns:bpmndi="http://www.omg.org/spec/BPMN/20100524/DI"
    xmlns:dc="http://www.omg.org/spec/DD/20100524/DC"
    xmlns:di="http://www.omg.org/spec/DD/20100524/DI"
    xmlns:xsi="http://www.w3.org/2001/XMLSchema-instance"
    id="Definitions_1"
    targetNamespace="http://bpmn.io/schema/bpmn">

A2. SWIMLANE ARCHITECTURE (MANDATORY)
──────────────────────────────────────
You MUST use Pools and Lanes to represent organizational boundaries.
You must wrap your process in a <bpmn:collaboration id="Collaboration_1"> element which contains the participant.

POOL — Represents the organisation. Use ONE expanded pool inside the collaboration: <bpmn:participant id="Pool_Main" name="Accounts Payable Process" processRef="Process_1" />
LANES — Represent roles/departments within the pool. Create a separate <bpmn:laneSet>
        and one <bpmn:lane> per unique responsible_role found in the process steps.
        Common AP lanes (adapt to actual roles in the data):
            • Vendor / Supplier
            • AP Clerk (Accounts Payable Team)
            • AP Manager / Finance Manager
            • Procurement / Purchasing
            • Controller / CFO
            • ERP System (for automated system tasks)

Each task MUST be referenced inside its owning lane via <bpmn:flowNodeRef>TaskId</bpmn:flowNodeRef>.
Lanes MUST have unique IDs: Lane_Vendor, Lane_APClerk, Lane_Manager, Lane_ERP, etc.

A3. TASK TYPE MAPPING (MANDATORY)
──────────────────────────────────
Map each process step to the correct BPMN 2.0 task type:
• <bpmn:userTask>     — Human performs the task manually (AP clerk keying data, manager approving)
• <bpmn:serviceTask>  — Automated system action (ERP auto-posting, email notification send)
• <bpmn:manualTask>   — Purely physical/paper-based step (printing, stamping, filing)
• <bpmn:sendTask>     — Sending a message/email to external party (sending remittance to vendor)
• <bpmn:receiveTask>  — Waiting to receive a response (waiting for vendor to re-submit)
• <bpmn:callActivity> — Invoking a sub-process (use for complex 3-way match sub-process)

All tasks must have: id, name attributes. Example:
    <bpmn:userTask id="Task_DataEntry" name="Key Invoice Data into ERP">

A4. GATEWAY RULES (CRITICAL — NEVER SKIP)
──────────────────────────────────────────
You MUST include ALL relevant gateway types. A straight linear flow is INCORRECT for any real AP process.

① EXCLUSIVE GATEWAY (XOR) — id prefix: GW_XOR_
    Use for: Yes/No decisions, single-path branches.
    Required in AP: Invoice Validation, 3-Way Match, Approval Decision, Duplicate Check.
    Example:
    <bpmn:exclusiveGateway id="GW_XOR_Validation" name="Invoice Valid?" />
    Outgoing flows MUST have a name: "Yes" / "No" / "Match" / "Mismatch"
    The "No" / exception path routes to an exception-handling task, NOT to an End Event directly.

② PARALLEL GATEWAY (AND) — id prefix: GW_AND_
    Use for: Tasks that happen simultaneously.
    Required in AP: When AP Clerk sends notification to Procurement AND updates the ERP log simultaneously.
    Example:
    <bpmn:parallelGateway id="GW_AND_Split" name="" />
    All outgoing branches MUST eventually merge at a closing <bpmn:parallelGateway id="GW_AND_Join">.

③ INCLUSIVE GATEWAY (OR) — id prefix: GW_OR_
    Use for: One or more conditional paths can be taken.
    Required in AP: When payment approval requires sign-off from Finance Manager OR Controller (amount-based).

④ EVENT-BASED GATEWAY — id prefix: GW_EVT_
    Use for: Process waits for one of several events.
    Required in AP: After sending discrepancy notice, wait for either vendor response OR internal resolution timeout.

A5. EVENT RULES (MANDATORY)
────────────────────────────
START EVENTS:
    • ONE <bpmn:startEvent> per process. Common AP start: "Invoice Received" (message start event).
    • Use <bpmn:messageEventDefinition> inside startEvent if triggered by an incoming invoice email.
    • id="StartEvent_InvoiceReceived"

END EVENTS (MULTIPLE — model all terminal states):
    • "Payment Completed" — Normal happy-path end.
    • "Invoice Rejected" — Terminated after vendor fails to correct.
    • "Duplicate Invoice Cancelled" — Terminated after duplicate detection.
    • "Escalated to Controller" — Error end event if approval chain escalates beyond normal.
    Use <bpmn:terminateEventDefinition> for error/abort ends, plain end event for normal completion.

INTERMEDIATE EVENTS:
    • Use <bpmn:intermediateCatchEvent> with <bpmn:timerEventDefinition> for SLA timeout scenarios
    (e.g., "Awaiting Vendor Response — 5 business days timeout" triggers escalation).
    • Use <bpmn:boundaryEvent> attached to tasks for exception boundary catches:
        - Attach to the "3-Way Match" callActivity: catch a mismatch boundary error event.
        - Attach to "Receive Vendor Response": catch a timer event (response overdue).

A6. EXCEPTION & LOOP-BACK PATHS (CRITICAL)
────────────────────────────────────────────
Every gateway that branches "No" / "Rejected" / "Mismatch" MUST route to one of:
a) An exception task (e.g., "Resolve Discrepancy", "Return Invoice to Vendor", "Escalate to Manager")
b) A loop-back sequence flow returning to an EARLIER task for re-work
    (e.g., after vendor corrects invoice, loop back to "Validate Invoice Data")
c) A terminate End Event if the exception is unrecoverable.

MANDATORY exception paths for AP:
    1. Invalid Invoice → "Return to Vendor for Correction" → [receiveTask: Wait for Corrected Invoice] → loop back to Validate
    2. 3-Way Match Fail → "Raise Discrepancy Query" → [sendTask: Notify Procurement/Vendor] → [Event-Based GW] → Wait for Response or Timeout
    3. Approval Rejected → "Send Rejection Notice to Vendor" → End: Invoice Rejected
    4. Duplicate Detected → "Flag & Discard Duplicate" → End: Duplicate Cancelled
    5. Payment Error → "Investigate Payment Failure" → [serviceTask: Retry Payment] → loop or escalate

A7. DATA OBJECTS (RECOMMENDED)
───────────────────────────────
Represent key documents as BPMN Data Objects connected to tasks via Data Associations:
• <bpmn:dataObjectReference id="DO_Invoice" name="Vendor Invoice" />
• <bpmn:dataObjectReference id="DO_PO" name="Purchase Order" />
• <bpmn:dataObjectReference id="DO_GRN" name="Goods Receipt Note" />
• <bpmn:dataObjectReference id="DO_PaymentVoucher" name="Payment Voucher" />
Connect with: <bpmn:dataInputAssociation> and <bpmn:dataOutputAssociation> on tasks.

A8. SEQUENCE FLOW RULES
────────────────────────
• EVERY flow MUST have a unique id (prefix: SF_) and sourceRef + targetRef.
• Conditional flows from gateways MUST have a name attribute (e.g., name="Match Successful").
• Default flows from XOR gateways MUST use the default attribute on the gateway and xsi:type on the flow.
• NEVER leave a task or event with zero incoming flows (except the Start Event).
• NEVER leave a task or event with zero outgoing flows (except End Events).
• Loop-back flows ARE allowed — they represent re-work cycles (vendor re-submission loops).

═══════════════════════════════════════════════════════════════════
PART B — BPMN DI (DIAGRAM INTERCHANGE) LAYOUT RULES
═══════════════════════════════════════════════════════════════════

B1. DI SECTION STRUCTURE (MANDATORY)
─────────────────────────────────────
Every XML file MUST include a <bpmndi:BPMNDiagram> section AFTER the process definition:

<bpmndi:BPMNDiagram id="BPMNDiagram_1">
    <bpmndi:BPMNPlane id="BPMNPlane_1" bpmnElement="Collaboration_1">
    <!-- BPMNShape for every task, gateway, event, pool, lane -->
    <!-- BPMNEdge for every sequenceFlow -->
    </bpmndi:BPMNPlane>
</bpmndi:BPMNDiagram>

B2. COORDINATE SYSTEM — UNIFORM 200px GRID (CRITICAL)
─────────────────────────────────────────────────────
THE SINGLE MOST IMPORTANT RULE: Every diagram MUST be SPACIOUS and READABLE.
Elements must NEVER overlap. Always leave generous whitespace between all nodes.

Layout direction:   LEFT-TO-RIGHT (horizontal flow)
Canvas origin:      x=160, y=80
Pool starts at:     x=160, y=80

LANE STRUCTURE & VERTICAL SPACING (y-axis):
    • Lane height is EXACTLY 200px (MANDATORY — do not use 240px or other values).
    • Lane label bar width (left side): 30px
    • Lane 1 (Vendor):     top y=80,   center y=180,  bottom y=280
    • Lane 2 (AP Clerk):   top y=280,  center y=380,  bottom y=480
    • Lane 3 (Manager):    top y=480,  center y=580,  bottom y=680
    • Lane 4 (ERP System): top y=680,  center y=780,  bottom y=880
    • Nodes are centered vertically within their lanes (y = lane_center_y).
    • Pool total height = number_of_lanes × 200

HORIZONTAL SPACING (x-axis):
    • Main happy path nodes are centered at x-coordinates spaced exactly 200px apart:
    - Start Event:  cx=260
    - Task 1:       cx=460
    - Gateway 1:    cx=660
    - Task 2:       cx=860
    - Gateway 2:    cx=1060
    - Task 3:       cx=1260
    - Gateway 3:    cx=1460
    - Task 4:       cx=1660
    - End Event:    cx=1860
    • Exception/branch tasks placed in a different lane use the EXACT SAME center x-coordinate as their source gateway.
    • Pool total width = (number_of_main_steps × 200) + 300 (for start/end padding)

Standard element dimensions (FIXED):
    Task (userTask, serviceTask, manualTask, sendTask, receiveTask):  width=120, height=80
    Gateway (all types):    width=50,  height=50
    Start Event / End Event: width=36,  height=36

TOP-LEFT BOUNDS FORMULAS (cx, cy are centers):
    • Task:    x = cx - 60,  y = cy - 40  (bounds: width="120" height="80")
    • Gateway: x = cx - 25,  y = cy - 25  (bounds: width="50" height="50")
    • Event:   x = cx - 18,  y = cy - 18  (bounds: width="36" height="36")

B3. LANE POSITIONING
─────────────────────
Pool shape (starts at x=160, y=80):
    <bpmndi:BPMNShape id="Pool_Main_di" bpmnElement="Pool_Main" isHorizontal="true">
    <dc:Bounds x="160" y="80" width="[CALCULATED_WIDTH]" height="[CALCULATED_HEIGHT]" />
    </bpmndi:BPMNShape>

Each lane shape starts at x=190 (pool x + 30) and height=200:
    <bpmndi:BPMNShape id="Lane_APClerk_di" bpmnElement="Lane_APClerk" isHorizontal="true">
    <dc:Bounds x="190" y="280" width="[pool_width - 30]" height="200" />
    </bpmndi:BPMNShape>

B4. SHAPE PLACEMENT — WORKED EXAMPLE (200px GRID)
──────────────────────────────────────────────────
For a process with:
- AP Clerk lane: center y=380 (bounds: y=280 to 480)
- Vendor lane:   center y=180 (bounds: y=80 to 280)
- Manager lane:  center y=580 (bounds: y=480 to 680)

Main happy-path shapes (AP Clerk lane, center y=380):
    Start Event:                 cx=260, cy=380 -> x="242" y="362"
    Task 1 (Data Entry):         cx=460, cy=380 -> x="400" y="340"
    GW 1 (Invoice Valid?):       cx=660, cy=380 -> x="635" y="355"
    Task 2 (3-Way Match):        cx=860, cy=380 -> x="800" y="340"
    GW 2 (Match OK?):            cx=1060,cy=380 -> x="1035"y="355"
    Task 3 (Approval):           cx=1260,cy=380 -> x="1200"y="340"
    GW 3 (Approved?):            cx=1460,cy=380 -> x="1435"y="355"
    Task 4 (Payment):            cx=1660,cy=380 -> x="1600"y="340"
    End Event (Payment Done):    cx=1860,cy=380 -> x="1842"y="362"

Exception shapes:
    - Return to Vendor (Vendor lane, cy=180): cx=660 -> x="600" y="140"
    - Receive Corrected (Vendor lane, cy=180): cx=860 -> x="800" y="140"
    - Resolve Discrepancy (Manager lane, cy=580): cx=1060 -> x="1000" y="540"
    - Rejection Notice (Vendor lane, cy=180): cx=1460 -> x="1400" y="140"
    - Rejection End Event (Vendor lane, cy=180): cx=1660 -> x="1642" y="162"

B5. EDGE (SEQUENCEFLOW) WAYPOINTS — WORKED EXAMPLE
───────────────────────────────────────────────────
Every sequence flow MUST route orthogonally (straight lines or L-shapes). NO DIAGONALS!

1. Straight Horizontal (e.g. Start Event to Task 1):
    <di:waypoint x="278" y="380" />   <!-- cx + 18 (right edge) -->
    <di:waypoint x="400" y="380" />   <!-- target x (left edge) -->

2. Straight Vertical Going UP (e.g. GW 1 to Return to Vendor):
    <di:waypoint x="660" y="355" />   <!-- gw cy - 25 (top edge) -->
    <di:waypoint x="660" y="220" />   <!-- task cy + 40 (bottom edge) -->

3. Straight Vertical Going DOWN (e.g. GW 2 to Resolve Discrepancy):
    <di:waypoint x="1060" y="405" />  <!-- gw cy + 25 (bottom edge) -->
    <di:waypoint x="1060" y="540" />  <!-- task cy - 40 (top edge) -->

4. Loop-back Going LEFT & DOWN (e.g. Receive Corrected to Task 1):
    Route UP near top of Vendor lane (y=110) then LEFT to x=460 and DOWN into Task 1:
    <di:waypoint x="860" y="140" />   <!-- task cy - 40 (top edge) -->
    <di:waypoint x="860" y="110" />   <!-- travel up to top of lane -->
    <di:waypoint x="460" y="110" />   <!-- travel left to target x -->
    <di:waypoint x="460" y="340" />   <!-- arrive at target top edge -->

5. Loop-back Going LEFT & UP (e.g. Resolve Discrepancy to Task 2):
    Route DOWN near bottom of Manager lane (y=650) then LEFT to x=860 and UP into Task 2:
    <di:waypoint x="1060" y="620" />  <!-- task cy + 40 (bottom edge) -->
    <di:waypoint x="1060" y="650" />  <!-- travel down to bottom of lane -->
    <di:waypoint x="860" y="650" />   <!-- travel left to target x -->
    <di:waypoint x="860" y="420" />   <!-- arrive at target bottom edge -->

B6. ANTI-OVERLAP RULES (STRICTLY ENFORCED)
────────────────────────────────────────────
1. No two shapes share overlapping bounding boxes.
2. Exception tasks placed off the main path MUST align vertically with their source gateways on the X axis.
3. Loop-back edges MUST route outside of active tasks by moving to the top or bottom lane margins (y=110 or y=650) before horizontal traversal.
4. Do not compress horizontal spacing below 200px between main-path node centers.
5. All edge sequence flows must enter and exit nodes at their exact centers vertically or horizontally.


═══════════════════════════════════════════════════════════════════
PART C — AP DOMAIN KNOWLEDGE (ENFORCE IN ALL DIAGRAMS)
═══════════════════════════════════════════════════════════════════

C1. MANDATORY AP PROCESS CHECKPOINTS
──────────────────────────────────────
Regardless of what steps are provided in the input, ALWAYS ensure the following
critical AP control points appear in the diagram. If a matching step exists in the
input data, use its name/description. If not, infer and add it:

① Duplicate Invoice Check (before data entry or after)
② Invoice-PO 2-Way or 3-Way Match Gateway (match vs mismatch branches)
③ Invoice Approval Gateway (approved vs rejected branches)
④ Payment Authorisation Gateway (within-limit vs above-limit → escalation)
⑤ Payment Execution & Confirmation

C2. ROLE-TO-LANE MAPPING
──────────────────────────
Vendor / Supplier            → Lane_Vendor
AP Clerk / AP Team / Finance Ops → Lane_APClerk
AP Manager / Finance Manager → Lane_Manager
Procurement / Buyer          → Lane_Procurement
Controller / CFO / Director  → Lane_Controller
ERP System / SAP / QuickBooks / Oracle / NetSuite → Lane_ERP
If a role doesn't cleanly map, create a new lane with an appropriate name.

C3. AUTOMATION HIGHLIGHTING
─────────────────────────────
Steps marked as is_manual=false in the input data MUST be modelled as <bpmn:serviceTask>.
Steps marked as is_manual=true MUST be modelled as <bpmn:userTask> or <bpmn:manualTask>.
High-ACS steps (automation candidate score > 7) SHOULD be placed in Lane_ERP or noted
as service tasks to visually indicate automation opportunity.

C4. REALISTIC AP EXCEPTION SCENARIOS TO INCLUDE
─────────────────────────────────────────────────
Always model AT MINIMUM these three exception sub-flows:
a) MISMATCH EXCEPTION: Invoice ≠ PO → Raise query → Vendor corrects → Re-validate
b) APPROVAL REJECTION: Manager rejects → Notify vendor → Close with rejection end event
c) OVERDUE PAYMENT: Timer boundary on payment task → Escalate to Controller

═══════════════════════════════════════════════════════════════════
PART D — QUALITY CHECKLIST (SELF-VALIDATE BEFORE OUTPUT)
═══════════════════════════════════════════════════════════════════

Before generating the final XML, mentally verify ALL of the following:
□ Every element has a globally unique ID (no duplicates across tasks, gateways, events, flows)
□ All namespace declarations are present in <bpmn:definitions>
□ Every lane references its tasks via <bpmn:flowNodeRef>
□ Every XOR gateway has exactly 2+ outgoing flows, each named
□ Every Parallel gateway has a matching closing parallel gateway
□ No task has zero incoming sequence flows (except Start Event)
□ No task has zero outgoing sequence flows (except End Events)
□ ALL Start Events, End Events, Tasks, Gateways, and Pools have a BPMNShape in the DI section
□ ALL Sequence Flows have a BPMNEdge in the DI section
□ Lane shapes have isHorizontal="true"
□ No element coordinates overlap (shapes do not visually collide)
□ Every exception path terminates at either a named End Event or a loop-back
□ The XML is well-formed (all tags properly opened and closed)

If ANY checklist item fails, fix it before outputting."""


def _generate_bpmn_xml(scored_steps: list) -> str:
    """Helper to generate BPMN XML from scored steps."""
    # Build a rich context block with ALL available step metadata
    steps_context = []
    for step in scored_steps:
        manual_flag = "MANUAL" if step.get("is_manual", True) else "AUTOMATED"
        acs = step.get("acs", step.get("automation_candidate_score", "N/A"))
        systems = ", ".join(step.get("systems_used", [])) or "N/A"
        steps_context.append(
            f"Step {step.get('step_number')} | [{manual_flag}] | Role: {step.get('responsible_role', 'AP Clerk')} "
            f"| Name: {step.get('name', '')} | Desc: {step.get('description', '')} "
            f"| Input: {step.get('input_document', 'N/A')} | Output: {step.get('output_document', 'N/A')} "
            f"| Systems: {systems} | Time: {step.get('estimated_time_minutes', '?')} min | ACS: {acs}"
        )

    # Derive unique roles for lane planning
    unique_roles = list(dict.fromkeys(
        step.get("responsible_role", "AP Clerk") for step in scored_steps
    ))

    prompt = f"""Generate a clean, spacious, industry-standard BPMN 2.0 XML diagram for the following
Accounts Payable process. READABILITY is the top priority — the diagram must be clean and not cramped.
Strictly follow every layout rule in your system prompt, especially Part B (spacing rules).

━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
PROCESS STEPS (use these as the backbone of the diagram)
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
{chr(10).join(steps_context)}

━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
ROLES IDENTIFIED (create one swimlane lane per role)
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
{chr(10).join(f'  • {{r}}' for r in unique_roles)}
Also include a "Vendor" lane for vendor-facing tasks and an "ERP System" lane for automated steps.
IMPORTANT: If there are already 4+ roles from the step data, do NOT add extra lanes — keep it to max 5 lanes total.

━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
DIAGRAM REQUIREMENTS
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

1. POOL & LANES:
- ONE expanded horizontal pool: name="Accounts Payable Process", id="Pool_Main"
- MAX 5 lanes (merge similar roles into one lane if needed to keep it clean)
- Each lane height: 200px (MANDATORY — do not use less)
- Each lane isHorizontal="true"

2. TASK TYPES:
- MANUAL steps → <bpmn:userTask> (human action)
- AUTOMATED steps → <bpmn:serviceTask> (system action)
- Vendor communication → <bpmn:sendTask> / <bpmn:receiveTask>

3. GATEWAYS & DECISION POINTS:
- Analyze the provided PROCESS STEPS and insert logical gateways where a decision or validation would naturally occur (e.g., "Data Valid?", "Approval Required?", "Match Successful?").
- For each gateway, create a "Yes / Success" path that continues the main flow.
- Create a "No / Failure" path that routes to a logical exception-handling task (e.g., "Resolve Discrepancy", "Return to Submitter", "Handle Error") and then either loops back or goes to a termination end event.
- Use <bpmn:exclusiveGateway> for XOR decisions.

4. EVENTS:
- Start Event: Create a descriptive start event based on the process context (e.g., "Process Triggered", "Request Received"). Place it in the lane of the first task.
- End Events: Create appropriate end events (e.g., "Process Complete", "Request Rejected").
- Place exception end events in the lane where the exception is handled.

5. LAYOUT — SPACING (READ CAREFULLY):
- Horizontal spacing between node CENTERS on the main happy path: exactly 200px
- Lane height: 200px each
- Main happy-path nodes stay in the AP Clerk lane (or relevant lane)
- Exception tasks for "No" gateway branches go in the SAME X-COLUMN as the gateway but in a LOWER or UPPER lane
- Place exception end events at X+200 from their exception task
- Example coordinate sequence for main path (assuming center y=380 for main lane):
    Start Event:  x=242, y=362  (center: 260, 380)
    Task 1:       x=400, y=340  (center: 460, 380)
    Gateway 1:    x=635, y=355  (center: 660, 380)
    Task 2:       x=800, y=340  (center: 860, 380)
    Gateway 2:    x=1035, y=355  (center: 1060, 380)
    Task 3:       x=1200, y=340  (center: 1260, 380)
    End Event:    x=1442, y=362  (center: 1460, 380)
- Exception tasks aligned vertically on the X-axis:
    Gateway 1 "No" → exception task at: x=600, y=140  (center: 660, 180 in a different lane)
    Gateway 2 "No" → exception task at: x=1000, y=540 (center: 1060, 580 in a different lane)
- DO NOT place any two nodes at the same x,y coordinates
- Total pool width = last_node_x + 150 (right padding) — minimum 1600px
- Total pool height = num_lanes × 200

6. EDGES - STRICT ROUTING RULES (NO DIAGONAL LINES ALLOWED):
CRITICAL: Edges MUST be horizontal OR vertical OR L-shaped (2 perpendicular segments).
NEVER draw a direct diagonal line from one lane to another.

a) Same-lane horizontal edge (2 waypoints, identical y):
    <di:waypoint x="278" y="380" />
    <di:waypoint x="400" y="380" />

b) Straight Vertical Going UP (e.g. Gateway to upper lane task):
    <di:waypoint x="660" y="355" />
    <di:waypoint x="660" y="220" />

c) Straight Vertical Going DOWN (e.g. Gateway to lower lane task):
    <di:waypoint x="1060" y="405" />
    <di:waypoint x="1060" y="540" />

d) Loop-back Going LEFT & DOWN:
    <di:waypoint x="860" y="140" />
    <di:waypoint x="860" y="110" />
    <di:waypoint x="460" y="110" />
    <di:waypoint x="460" y="340" />

e) Loop-back Going LEFT & UP:
    <di:waypoint x="1060" y="620" />
    <di:waypoint x="1060" y="650" />
    <di:waypoint x="860" y="650" />
    <di:waypoint x="860" y="420" />

7. IDs: All IDs must be unique and descriptive based on the actual step names (e.g., Task_ExtractData, GW_XOR_Validation, Task_ResolveError).

8. OUTPUT: ONLY the raw BPMN 2.0 XML inside ```xml ... ```. No prose. No comments outside the XML.
"""

    xml_result = call_gemini_text(
        prompt=prompt,
        system_prompt=BPMN_SYSTEM_PROMPT,
        model="gemini-3.5-flash"
    )
    
    # Clean the output if it has markdown formatting
    cleaned_xml = xml_result.strip()
    if cleaned_xml.startswith("```xml"):
        cleaned_xml = cleaned_xml[6:]
    if cleaned_xml.startswith("```"):
        cleaned_xml = cleaned_xml[3:]
    if cleaned_xml.endswith("```"):
        cleaned_xml = cleaned_xml[:-3]
        
    return cleaned_xml.strip()


@app.post("/api/flowchart/generate")
async def generate_flowchart(payload: dict = Body(...)):
    """Generate intelligent BPMN 2.0 XML diagram."""
    scored_steps = payload.get("scored_steps", [])
    if not scored_steps:
        raise HTTPException(status_code=400, detail="scored_steps is required")

    try:
        cleaned_xml = _generate_bpmn_xml(scored_steps)
        return {"status": "success", "data": {"xml": cleaned_xml}}
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


@app.post("/api/flowchart/modify")
async def modify_flowchart(payload: dict = Body(...)):
    """Modify an existing BPMN diagram based on user prompt."""
    current_xml = payload.get("xml")
    user_prompt = payload.get("prompt")
    
    if not current_xml or not user_prompt:
        raise HTTPException(status_code=400, detail="Both 'xml' and 'prompt' are required")

    prompt = f"""You are an expert BPMN 2.0 process architect modifying an existing AP process diagram.

USER INSTRUCTION: {user_prompt}

CURRENT BPMN 2.0 XML:
```xml
{current_xml}
```

MODIFICATION RULES:
1. Apply ONLY the changes requested by the user instruction above.
2. Preserve ALL existing elements, IDs, namespaces, and DI coordinates that are not affected by the change.
3. If adding new elements (tasks, gateways, events), assign them unique IDs following the existing naming convention.
4. If adding new tasks, also add their BPMNShape entries in the DI section with non-overlapping coordinates.
5. If adding new sequence flows, also add their BPMNEdge entries with correct waypoints.
6. If removing elements, also remove their corresponding BPMNShape/BPMNEdge and any orphaned sequence flows.
7. Ensure the result passes the quality checklist:
- No element has 0 incoming flows (except Start Event)
- No element has 0 outgoing flows (except End Events)
- All gateway branches are named
- All IDs remain globally unique
8. OUTPUT ONLY the complete, updated BPMN 2.0 XML inside a ```xml code block. No other text.
"""

    try:
        xml_result = call_gemini_text(
            prompt=prompt,
            system_prompt=BPMN_SYSTEM_PROMPT,
            model="gemini-3.5-flash"
        )
        
        cleaned_xml = xml_result.strip()
        if cleaned_xml.startswith("```xml"):
            cleaned_xml = cleaned_xml[6:]
        if cleaned_xml.startswith("```"):
            cleaned_xml = cleaned_xml[3:]
        if cleaned_xml.endswith("```"):
            cleaned_xml = cleaned_xml[:-3]
            
        cleaned_xml = cleaned_xml.strip()
            
        return {"status": "success", "data": {"xml": cleaned_xml}}
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

@app.post("/api/chat")
async def chat_with_agent(payload: dict = Body(...), user_id: str = Depends(get_current_user)):
    """Ask queries about or modify the generated report."""
    import json
    from pydantic import BaseModel
    
    class ChatIntent(BaseModel):
        intent: str
    
    context_data = payload.get("context_data")
    messages = payload.get("messages", [])
    pending_doc_edits = payload.get("pending_doc_edits", [])
    
    if not messages:
        raise HTTPException(status_code=400, detail="Messages are required")

    # Format conversation history (limit to last 15 messages to save tokens/costs)
    history_str = ""
    recent_messages = messages[:-1][-15:] if len(messages) > 1 else []
    for msg in recent_messages:
        role = msg.get("role", "User").upper()
        history_str += f"{role}: {msg.get('content')}\n"
        
    latest_query = messages[-1].get("content")

    # --- ROUTER LOGIC --- (always run, even when pending_doc_edits exist)
    # pending_doc_edits = highlighted text comments queued by the user.
    # The user's typed message determines intent:
    #   "explain this" / "what does this mean" → ask mode (explain in chat)
    #   "update this" / "rewrite" / no message  → update_doc mode (edit the document)
    
    # Build pending edits summary for router context
    edits_summary = ""
    if pending_doc_edits:
        edits_summary = "\n\nPENDING DOCUMENT EDITS (highlighted text comments queued by user):\n"
        for i, edit in enumerate(pending_doc_edits):
            edits_summary += f"{i+1}. Selected text: \"{edit.get('selectedText', '')}\"\n   Comment/instruction: {edit.get('comment', '')}\n"

    # Default: if there are pending edits but the user typed nothing, that means apply edits
    default_mode = "update_doc" if pending_doc_edits else "ask"

    # If the user just hit Send without typing in the main input, the frontend sends a placeholder.
    # We MUST replace the placeholder with the actual comments BEFORE the router sees it, 
    # otherwise the router gets confused by the words "apply the pending edits".
    if latest_query.strip() == "Please apply the pending document edits." and pending_doc_edits:
        comments = [e.get('comment', '').strip() for e in pending_doc_edits if e.get('comment', '').strip()]
        if comments:
            latest_query = " | ".join(comments)
        else:
            latest_query = "Please update the document."

    # Smart default fallback in case the AI router fails
    # If the user submitted via the "Apply pending edits" button with multiple edits, it's a batch operation.
    original_query = messages[-1].get("content", "").strip()
    is_batch_apply = (original_query == "Please apply the pending document edits." and len(pending_doc_edits) > 1)

    is_question = False
    if not is_batch_apply:
        # Only use simple keyword matching to force "ask" mode if it's not a multi-edit batch
        is_question = any(word in latest_query.lower() for word in ["explain", "what", "how", "summarize", "why", "tell me"])
    
    if is_question:
        default_mode = "ask"
    else:
        default_mode = "update_doc" if pending_doc_edits else "ask"

    # Bypass the router completely for obvious questions to ensure 100% reliability
    if is_question:
        mode = "ask"
    else:
        router_prompt = f"""You are an intelligent router for an AP Automation Agent.
The user is interacting with a document (PDD or SDD) or dashboard. They may have highlighted text and left a comment. Their typed message tells you what they actually want.

LATEST QUERY: "{latest_query}"
CONVERSATION HISTORY (for context):
{history_str}{edits_summary}

Determine the user's intent:
- "ask": The user is asking a question, wants an explanation, wants to understand the content, or wants the AI to answer in the chat. Examples: "explain this", "what does this mean", "why is this here", "tell me more".
- "edit": The user wants to change, update, fix, or add something to the overall assessment data/report/dashboard.
- "update_doc": The user wants the AI to edit/update/rewrite the PDD or SDD document itself. Examples: "update this", "rewrite", "fix this", "make this shorter", "apply edits". Also use this if they ask to restore a SPECIFIC deleted section.
- "revert": The user explicitly wants to undo a generic LAST change made to the project, dashboard, or document (like hitting Ctrl+Z). Examples: "revert that", "undo my last edit", "put it back the way it was", "undo".

IMPORTANT ROUTING RULES:
1. If the user asks to revert a SPECIFIC named section (like "revert section 1.1"), map it to "update_doc", NOT "revert".
2. If there are MULTIPLE pending document edits containing a mix of editing instructions (e.g., "remove this", "add a summary") and questions (e.g., "explain this"), you MUST prioritize "update_doc" (or "edit"). The AI will apply the edits and answer the questions simultaneously.
3. Only use "ask" if the SOLE intent is to ask a question or get an explanation without modifying the document.
"""
        try:
            router_result = call_gemini_structured(
                prompt=router_prompt,
                system_prompt="You are a strict JSON intent router.",
                response_schema=ChatIntent
            )
            mode = router_result.get("intent", default_mode).lower()
            if mode not in ["ask", "edit", "update_doc", "revert"]:
                mode = default_mode
        except Exception as e:
            print(f"Router error: {e}")
            mode = default_mode  # fallback

    # Strip out bpmn_xml and chat_history — they're huge and irrelevant for the chat AI
    if context_data:
        context_data = {k: v for k, v in context_data.items() if k not in ('bpmn_xml', 'chat_history')}
    context_str = json.dumps(context_data, indent=2)

    # If the router mapped this to "ask" mode but the query is just the auto-placeholder,
    # replace the query with the user's actual comment so the Chat AI doesn't get confused and try to "apply edits".
    if mode == "ask" and latest_query.strip() == "Please apply the pending document edits." and pending_doc_edits:
        # Take all comments and join them, or use a default if none exist
        comments = [e.get('comment', '').strip() for e in pending_doc_edits if e.get('comment', '').strip()]
        if comments:
            latest_query = " | ".join(comments)
        else:
            latest_query = "Please explain the highlighted text."

    # Handle doc_type resolution and prevent mixed batch edits
    if pending_doc_edits:
        doc_types = set(e.get("docType") for e in pending_doc_edits if e.get("docType"))
        
        has_dashboard = "dashboard" in doc_types
        has_doc = "pdd" in doc_types or "sdd" in doc_types
        
        if has_dashboard and has_doc:
            return {
                "status": "success", 
                "data": {
                    "response_text": "⚠️ **Mixed Edits Detected:** You have selected comments from both the Dashboard and a Document (PDD/SDD). Because Dashboard updates and Document rewrites use different AI pipelines, please apply your Dashboard edits and Document edits separately."
                }
            }
            
        if has_dashboard and mode == "update_doc":
            mode = "edit"

    # Build highlighted context for both 'ask' and 'edit' modes
    highlighted_context = ""
    if pending_doc_edits:
        highlighted_context = "\n---\nHIGHLIGHTED TEXT FROM THE UI (the user is referring to this):\n"
        for i, edit in enumerate(pending_doc_edits):
            sel = edit.get('selectedText', '').strip()
            comment = edit.get('comment', '').strip()
            if sel:
                highlighted_context += f"\nSelection {i+1}:\n{sel}\n"
            if comment:
                highlighted_context += f"User's note: {comment}\n"

    if mode == "ask":
        prompt = f"""You are an expert Accounts Payable Automation Consultant in READ-ONLY Mode.
The user has uploaded an AP process, and an AI agent has generated an analysis report based on it.

Here is the full analysis context (JSON format):
{context_str}
{highlighted_context}
---
CONVERSATION HISTORY:
{history_str}

---
LATEST USER QUERY:
{latest_query}

Please answer the user's query comprehensively based on the analysis context and any highlighted text shown above.
IMPORTANT: You are in "Ask Mode". Explain the content directly in your chat response — do NOT modify the document. Use Markdown for formatting.
"""
        try:
            response_text = call_gemini_text(
                prompt=prompt,
                system_prompt="You are an expert AP Automation Consultant. You provide clear, professional, and insightful advice based on the provided JSON analysis.",
                model="gemini-3.5-flash"
            )
            return {"status": "success", "data": {"response_text": response_text}}
        except Exception as e:
            raise HTTPException(status_code=500, detail=str(e))

            
    elif mode == "edit":
        prompt = f"""You are an expert Accounts Payable Automation Consultant in AGENT EDITING Mode.
The user has uploaded an AP process, and an AI agent has generated an analysis report based on it.

Here is the full analysis context (JSON format):
{context_str}
{highlighted_context}
---
CONVERSATION HISTORY:
{history_str}

---
LATEST USER QUERY:
{latest_query}

You have permission to edit the user's data (the context above). 
If the user asks you to modify the report, add a summary, or change any scores/steps, you must modify the corresponding structured data AND the `markdown_report` to keep everything perfectly in sync.
IMPORTANT: When you add, delete, or move a section in the `markdown_report`, you MUST recalculate and update all section numbers (indices) in the entire report so they remain sequential (e.g., if section 1.1 is deleted, the old 1.2 becomes the new 1.1).
OUTPUT INSTRUCTIONS:
You must output a single, raw JSON object inside a ```json codeblock. The JSON must have EXACTLY these two keys:
"response_text": A string explaining to the user what you changed.
"updated_context_data": The FULL, completely rewritten context_data object. If no changes are needed, return the original context_data.

Output ONLY the JSON codeblock. No other text.
"""

        try:
            parsed_data = call_gemini_structured(
                prompt=prompt,
                system_prompt="You are an expert AP Automation Agent capable of rewriting structured JSON and markdown reports based on user instructions.",
                model="gemini-3.5-flash"
            )
            
            # Extract updated context data and save it to MongoDB
            updated_context = parsed_data.get("updated_context_data")
            assessment_id = context_data.get("id")
            if updated_context and assessment_id:
                # Ensure the 'id' field is preserved no matter what the AI did!
                updated_context["id"] = assessment_id
                if "_id" in context_data:
                    updated_context["_id"] = context_data["_id"]
                    
                # Save the new state back to the database
                update_assessment_data(assessment_id, user_id, updated_context)
            
            return {"status": "success", "data": parsed_data}
        except Exception as e:
            raise HTTPException(status_code=500, detail=str(e))
            
    elif mode == "revert":
        assessment_id = context_data.get("id")
        if not assessment_id:
            raise HTTPException(status_code=400, detail="assessment_id missing from context_data")

        restored_doc = revert_global_assessment(assessment_id, user_id)
        if restored_doc:
            response_text = "✅ **Success:** I have perfectly undone your last change and restored the entire project from the database backups!"
            return {
                "status": "success", 
                "data": {
                    "response_text": response_text,
                    "updated_context_data": restored_doc,
                    "updated_document": {
                        "type": "pdd", # The frontend will refresh both automatically from updated_context_data anyway, but we return a generic structure
                        "content": restored_doc.get("documents", {}).get("pdd", "")
                    }
                }
            }
        else:
            return {"status": "success", "data": {"response_text": "⚠️ **Error:** I could not revert your changes. There is no previous backup available in the Undo memory stack."}}

    elif mode == "update_doc":
        assessment_id = context_data.get("id")
        if not assessment_id:
            raise HTTPException(status_code=400, detail="assessment_id missing from context_data")

        # Determine doc_type from pending edits or prompt
        doc_type = "pdd"
        if pending_doc_edits and pending_doc_edits[0].get("docType"):
            doc_type = pending_doc_edits[0].get("docType")
        elif "sdd" in latest_query.lower():
            doc_type = "sdd"
            
        assessment = get_assessment(assessment_id, user_id)
        if not assessment:
            raise HTTPException(status_code=404, detail="Assessment not found")
            
        documents = assessment.get("documents", {})
        current_doc = documents.get(doc_type)
        if not current_doc:
            return {"status": "success", "data": {"response_text": f"Error: No {doc_type.upper()} exists yet. Please generate one first."}}
            
        doc_label = "Process Design Document (PDD)" if doc_type == "pdd" else "Solution Design Document (SDD)"

        bpmn_xml = assessment.get("bpmn_xml", "")
        bpmn_block = ""
        if bpmn_xml:
            bpmn_block = f"```bpmn\n{bpmn_xml}\n```"
            import re
            pattern = r"```bpmn\s+[\s\S]*?```"
            current_doc_for_gemini = re.sub(pattern, "{{BPMN_DIAGRAM}}", current_doc)
        else:
            current_doc_for_gemini = current_doc

        # ── Split pending edits into: doc edits vs. questions/explanations ──────
        QUESTION_KEYWORDS = ["explain", "what", "why", "how", "describe", "tell me", "summarize", "what does", "what is"]

        doc_edit_comments  = []  # edits to apply to the document
        explain_comments   = []  # questions to answer in the chat

        for edit in pending_doc_edits:
            comment_lower = edit.get("comment", "").lower().strip()
            if any(kw in comment_lower for kw in QUESTION_KEYWORDS):
                explain_comments.append(edit)
            else:
                doc_edit_comments.append(edit)

        # ── Build instruction string only from doc edits ──────────────────────
        auto_placeholder = "Please apply the pending document edits."
        if latest_query and latest_query.strip() != auto_placeholder:
            instruction = f"User Instruction: {latest_query}\n\n"
        else:
            instruction = ""

        if doc_edit_comments:
            instruction += "Please apply the following specific edits:\n"
            for i, edit in enumerate(doc_edit_comments):
                instruction += f"{i+1}. For the text: \"{edit.get('selectedText')}\"\n   Edit requested: {edit.get('comment')}\n\n"

        # ── If there are no actual doc edits, nothing to write ────────────────
        has_doc_edits = bool(doc_edit_comments or (latest_query and latest_query.strip() != auto_placeholder))

        updated_content = None
        doc_response_text = ""
        explanation_text = ""

        try:
            # ── Step 1: Apply document edits (if any) ─────────────────────────
            if has_doc_edits:
                prompt = f"""You are a senior document specialist. Below is the current {doc_label}.

Here is the underlying structured JSON data for the project (in case you need to regenerate deleted sections or reference raw data):
{context_str}

---
CONVERSATION HISTORY (Use this to understand contextual commands like "revert my last change"):
{history_str}

---
The user wants you to make the following change to the document:
{instruction}

Apply the requested change to the document based on the user's instructions and the conversation history. Return the COMPLETE updated document with ALL sections intact — do not omit or summarize any sections that were not changed. Only modify the sections that the user's instruction/selection refers to, EXCEPT for section numbering: if you add, delete, or move a section, you MUST recalculate and update all section numbers (indices) in the entire document so they remain sequential (e.g., if section 1.1 is deleted, the old 1.2 becomes the new 1.1). Keep the same professional markdown formatting.
IMPORTANT: Do NOT modify the placeholder '{{BPMN_DIAGRAM}}'. Leave it exactly as it is.

Current document:
{current_doc_for_gemini}
"""
                updated_content = call_gemini_text(
                    prompt=prompt,
                    system_prompt="You are a senior document specialist. Apply the requested changes carefully. Return the complete updated document — never truncate or summarize unchanged sections.",
                    model="gemini-3.5-flash"
                )

                # Re-inject BPMN block
                if bpmn_xml and bpmn_block:
                    if "{{BPMN_DIAGRAM}}" in updated_content:
                        updated_content = updated_content.replace("{{BPMN_DIAGRAM}}", bpmn_block)
                    elif "{BPMN_DIAGRAM}" in updated_content:
                        updated_content = updated_content.replace("{BPMN_DIAGRAM}", bpmn_block)
                    elif "[[BPMN_DIAGRAM]]" in updated_content:
                        updated_content = updated_content.replace("[[BPMN_DIAGRAM]]", bpmn_block)

                update_assessment_document(assessment_id, user_id, doc_type, updated_content)
                doc_response_text = f"✅ I have successfully applied **{len(doc_edit_comments)} edit(s)** to the {doc_type.upper()}."

            # ── Step 2: Answer any "explain" questions in the chat ─────────────
            if explain_comments:
                explain_parts = "\n".join(
                    f"- Selected text: \"{e.get('selectedText', '')}\"\n  Question: {e.get('comment', '')}"
                    for e in explain_comments
                )
                explain_prompt = f"""You are an expert AP Automation Consultant.
The user has highlighted sections in a {doc_label} and wants you to explain them.

Here is the full project context:
{context_str}

---
Highlighted sections and questions from the user:
{explain_parts}

---
Please provide a clear, concise, professional explanation for each question. Use Markdown formatting with headers for each question.
"""
                explanation_text = call_gemini_text(
                    prompt=explain_prompt,
                    system_prompt="You are an expert AP Automation Consultant. Provide clear, insightful explanations.",
                    model="gemini-3.5-flash"
                )

            # ── Step 3: Combine response ──────────────────────────────────────
            combined_response = "\n\n---\n\n".join(filter(None, [doc_response_text, explanation_text]))
            if not combined_response:
                combined_response = "No edits or questions were provided."

            result_data = {"response_text": combined_response}
            if updated_content:
                result_data["updated_document"] = {
                    "type": doc_type,
                    "content": updated_content
                }

            return {"status": "success", "data": result_data}

        except Exception as e:
            raise HTTPException(status_code=500, detail=str(e))

    else:
        raise HTTPException(status_code=400, detail="Invalid mode specified")


# Restart trigger


# ═══════════════════════════════════════════════════════════════════
#  LIVE CHAT — Simple text-based AP Discovery Conversation
#  Uses GEMINI_LIVE_API key exclusively. Browser handles voice I/O.
# ═══════════════════════════════════════════════════════════════════

LIVE_CHAT_BASE_INSTRUCTIONS = """You are Aria — a warm, sharp, and genuinely curious business process consultant having a real-time voice conversation with a client. Your personality is like a brilliant colleague who actually enjoys diving into the messy details of how work gets done. You're friendly, encouraging, and make the client feel heard at every step.

Your mission: gather a COMPLETE picture of their business process so the system can generate a thorough, industry-grade automation analysis report. Do not generate the report yourself — your only job is discovery.

━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
REAL-TIME CONVERSATION RULES (NON-NEGOTIABLE)
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

STRICT SCOPE LIMITATION:
- You are strictly an Accounts Payable (AP) process discovery and automation consultant.
- Do NOT help with math, coding, homework, general knowledge, cheating, or any out-of-scope topics.
- Normal greetings, pleasantries, and questions about how you are doing are fully allowed.
- If the user asks or talks about anything out-of-scope, say: "I can't help with that. Let's focus on mapping your AP process."

LANGUAGE (SET BY SESSION CONFIG — SEE BELOW):
- Follow the exact language instruction specified below.

RESPONSE LENGTH:
- Keep every response under 2 sentences and ideally under 20 words.
- Long responses cause audio lag. Brevity is kindness here.

OPENING:
- Do NOT speak first on connection. Wait silently for the user to speak.

ONE QUESTION AT A TIME:
- Ask exactly ONE question per turn. Never stack questions.
- Acknowledge their answer in 2-4 words ("Got it.", "That makes sense.", "Oh interesting.") then ask your next question.

ECHO HANDLING:
- If the user input seems to repeat your own question back to you, ignore the repeat and ask them to clarify or move to the next question.

INTERRUPTIONS:
- If the user shifts topics or interrupts, immediately drop your current thread and follow their lead.

TONE — BE WARM AND HUMAN:
- Use contractions: "I'll", "let's", "that's", "we've", "doesn't", "you're"
- Be encouraging when they share pain points: "Ugh, yeah, that's a really common bottleneck."
- Show genuine curiosity: "Oh that's interesting — tell me more about that step."

INTERNAL TRACKING (SILENT):
- Keep a mental checklist of which of the domains below you have SOLID answers for.
- CRITICAL: NEVER speak your internal checklist, rules, or reasoning out loud!
- DO NOT NARRATE YOUR PROCESS: Never say things like "Initiating Discovery Phase" or "Moving to the next domain". Just ask the natural conversational question directly.

━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
HOW TO END THE CONVERSATION
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

Phase A — READINESS CHECK:
Before moving to the summary, verify ALL required domains are confirmed. If ANY are missing, keep asking.

Phase B — SUMMARIZE EVERYTHING (AND STOP):
Once all gates are green, deliver a full, detailed analysis summary of everything you understood from the user's answers. (Note: It is OK to exceed the 2-sentence length limit for this final summary).

"Alright, I think I've got a solid picture now! Here is a detailed summary of your process: [Provide a comprehensive, detailed summary covering the process name, key steps, main systems, roles involved, volume/time, and top pain points]. Does that sound right, or is there anything you'd like to add or correct?"

CRITICAL RULE FOR PHASE B: You MUST STOP speaking immediately after asking this question. DO NOT append [DISCOVERY_COMPLETE] to this summary. You MUST wait for the user to reply to your summary first!

Phase C — FINAL HANDOFF:
Only execute this phase AFTER the user has replied to your Phase B summary. Listen to their response:

-> If they say NO / nothing more / sounds good / that's all / we're done (or anything indicating completion):
    Respond with EXACTLY this, nothing more, nothing less:
    "Perfect — I'm sending this off to generate your full analysis report now. [DISCOVERY_COMPLETE]"

-> If they say YES / actually / wait / one more thing (or want to add/correct something):
    Acknowledge briefly ("Of course, go ahead."), collect the addition, then loop back to Phase B.

CRITICAL MARKER RULES:
- The exact text [DISCOVERY_COMPLETE] must appear at the VERY END of your response text, and ONLY in Phase C.
- NEVER output [DISCOVERY_COMPLETE] in the same turn as your Phase B summary.
- Do NOT add any text after [DISCOVERY_COMPLETE].
- Do NOT use [DISCOVERY_COMPLETE] at any other point in the conversation.
- Do NOT generate or summarize the analysis report yourself — your job ends at discovery."""


LIVE_CHAT_DEFAULT_CHECKLIST = """━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
INFORMATION YOU MUST COLLECT (all 8 domains)
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

1. PROCESS IDENTITY
- What is this process called / what business function does it serve?
- Is it triggered by an event (email, form, approval) or run on a schedule?
- Where does the process start and where does it officially end?

2. STEP-BY-STEP WORKFLOW
- Walk through each step from trigger to completion (aim for 5–10 distinct steps)
- For each step: What action is taken? Who does it? What input is needed? What output is produced?
- Are there any conditional branches, approvals, or exception paths?

3. SYSTEMS & TOOLS
- Every software, platform, ERP, CRM, database, or spreadsheet touched during the process
- Are any of these systems integrated with each other, or is data moved manually between them?
- Are there any legacy systems, custom-built tools, or paper-based steps?

4. PEOPLE & ROLES
- Who is involved at each step? (job titles / team names)
- How many people (total FTE) are dedicated to THIS process? (Needed to reconcile the ROI's freed-capacity figure.)
- Is there a supervisor, approver, or quality-checker in the loop?

5. VOLUME & TIME
- How many times is this process run per day / week / month? (Pin down a concrete number you can convert to an ANNUAL volume — this directly drives ROI.)
- How long does the full end-to-end process take per instance?
- Which specific step takes the longest? Which is most frequent?

6. PAIN POINTS & ERRORS
- Where do mistakes most often happen? What kind of errors?
- Roughly what % of transactions hit an exception / don't go straight-through? (Get a number — it drives the realistic time-saved and ROI.)
- Where do delays or bottlenecks occur?
- What causes rework, escalations, or complaints?
- Any compliance, audit, or regulatory concerns?

7. COST & BUSINESS IMPACT
- Which CURRENCY are we working in? (e.g. USD, GBP, EUR — every figure depends on this.)
- Approximate loaded hourly rate (or salary band) of the staff doing this work
- Cost of errors or delays when they happen (financial, reputational, or operational)
- Is there a peak season or time pressure that makes this process more critical?

8. EXISTING AUTOMATION & CONSTRAINTS
- Is anything already partially automated? What tools are in place?
- Have there been past automation attempts? What happened?
- Any hard constraints: budget, IT policy, data security, integrations that can't change?

━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
FIVE CRITICAL FACTS FOR ACCURATE ROI (GET A CONCRETE NUMBER FOR EACH)
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
The report's ROI and payback are computed from these five facts. Do not wrap up the call until you have a specific value (or an explicit "I don't know") for each:
1. Currency (USD / GBP / EUR / etc.)
2. Loaded hourly labour rate of the AP staff
3. Transaction/invoice volume you can convert to a yearly figure
4. Total FTE (headcount) working on this process
5. Exception rate — the % of transactions that need manual handling
If the client genuinely doesn't know one, acknowledge it and move on — the report will flag it as an assumption rather than invent a number.

━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
REAL-TIME CONVERSATION RULES (NON-NEGOTIABLE)
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

STRICT SCOPE LIMITATION:
- You are strictly an Accounts Payable (AP) process discovery and automation consultant.
- Do NOT help with math, coding, homework, general knowledge, cheating, or any out-of-scope topics.
- Normal greetings, pleasantries, and questions about how you are doing (e.g., "Hello", "Hi", "How are you?", "Greet me first") are fully allowed and should be answered warmly and briefly (e.g., "Hi there! I'm doing great, thanks for asking. How can I help you today?").
- If the user asks or talks about anything out-of-scope (other than greetings and pleasantries), you must say: "I can't help with that. Let's focus on mapping your AP process."

LANGUAGE (CRITICAL):
- You MUST always respond in the exact same language the user is speaking.
- If the user speaks English, you MUST strictly respond in English. Do NOT switch to Hindi or any other language unless the user speaks it first.

RESPONSE LENGTH:
- Keep every response under 2 sentences and ideally under 20 words.
- Long responses cause audio lag. Brevity is kindness here.

OPENING:
- Do NOT speak first on connection. Wait silently for the user to speak.

ONE QUESTION AT A TIME:
- Ask exactly ONE question per turn. Never stack questions.
- Acknowledge their answer in 2–4 words ("Got it.", "That makes sense.", "Oh interesting.", "Love that detail.") then ask your next question.

ECHO HANDLING:
- If the user input seems to repeat your own question back to you, ignore the repeat and ask them to clarify or move to the next question.

INTERRUPTIONS:
- If the user shifts topics or interrupts, immediately drop your current thread and follow their lead. Do not try to finish your previous question.

TONE — BE WARM AND HUMAN:
- Use contractions: "I'll", "let's", "that's", "we've", "doesn't", "you're"
- Be encouraging when they share pain points: "Ugh, yeah, that's a really common bottleneck."
- Show genuine curiosity: "Oh that's interesting — tell me more about that step."
- Avoid stiff corporate language. Sound like a smart friend, not a survey form.

LANGUAGE:
- Follow the strict language instruction specified in your session configuration below.

INTERNAL TRACKING (SILENT):
- Keep a mental checklist of which of the 8 domains above you have SOLID answers for.
- CRITICAL: NEVER speak your internal checklist, rules, or reasoning out loud! Because you are a voice AI, whatever text you generate is spoken directly to the user. Do not say things like "My internal state is..." or "I have registered your input". Keep your tracking completely invisible and silent.
- DO NOT NARRATE YOUR PROCESS: Never say things like "Initiating Discovery Phase", "Moving to the next domain", or "I am now focusing on...". Just ask the natural conversational question directly.

━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
HOW TO END THE CONVERSATION
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

Phase A — READINESS CHECK:
Before moving to the summary, verify ALL of the following are confirmed:
✅ Process name and trigger identified
✅ At least 5 distinct process steps described
✅ At least 2 systems/tools named
✅ Roles/headcount confirmed
✅ Volume AND time estimates for at least one step
✅ At least 2 pain points or error types
✅ Staff cost or business impact mentioned
✅ Existing automation level confirmed (even if "none")

If ANY of these are missing, keep asking. Do not proceed to the summary prematurely.

Phase B — SUMMARIZE EVERYTHING:
Once all 8 gates are green, deliver a full, detailed analysis summary of everything you understood from the user's answers. (Note: It is OK to exceed the 2-sentence length limit for this final summary).

"Alright, I think I've got a solid picture now! Here is a detailed summary of your process: [Provide a comprehensive, detailed summary covering the process name, key steps, main systems, roles involved, volume/time, and top pain points]. Does that sound right, or is there anything you'd like to add or correct before I hand this off for the final report generation?"

Phase C — FINAL HANDOFF:
Listen to their response:

→ If they say NO / nothing more / sounds good / that's all / we're done (or anything indicating completion):
    Respond with EXACTLY this, nothing more, nothing less:
    "Perfect — I'm sending this off to generate your full analysis report now. [DISCOVERY_COMPLETE]"

→ If they say YES / actually / wait / one more thing (or want to add/correct something):
    Acknowledge briefly ("Of course, go ahead."), collect the addition, then loop back to Phase B.

CRITICAL MARKER RULES:
- The exact text [DISCOVERY_COMPLETE] must appear at the VERY END of your response text.
- Do NOT add any text after [DISCOVERY_COMPLETE].
- Do NOT use [DISCOVERY_COMPLETE] at any other point in the conversation.
- Do NOT generate or summarize the analysis report yourself — your job ends at discovery."""


class LiveChatMessage(BaseModel):
    role: str  # 'user' or 'assistant'
    content: str


class LiveChatRequest(BaseModel):
    messages: List[LiveChatMessage]


@app.post("/api/analyze/live-chat")
async def live_chat(
    payload: LiveChatRequest,
    user_id: str = Depends(get_current_user)
):
    """Single text-based live chat endpoint for AP discovery.
    Uses GEMINI_LIVE_API key exclusively. Browser handles all voice I/O.
    Returns text response only — no TTS, no audio processing."""
    try:
        messages = payload.messages

        # Build conversation context for Gemini
        conversation_text = ""
        for msg in messages:
            role_label = "Client" if msg.role == "user" else "Consultant"
            conversation_text += f"{role_label}: {msg.content}\n\n"

        if not messages:
            # Greeting prompt — keep it very brief for fast response
            prompt = "Greet the client briefly (1 sentence) and ask them to describe how their AP process starts (1 sentence)."
        else:
            prompt = f"""Continue this AP discovery conversation. You are the Consultant.

{conversation_text}

Consultant:"""

        response_text = call_gemini_text(
            prompt=prompt,
            system_prompt=LIVE_CHAT_BASE_INSTRUCTIONS + "\n" + LIVE_CHAT_DEFAULT_CHECKLIST,
            model="gemini-2.0-flash",
            use_live_key=True
        )

        # Check if AI signaled satisfaction
        satisfied = "[DISCOVERY_COMPLETE]" in response_text
        clean_response = response_text.replace("[DISCOVERY_COMPLETE]", "").strip()

        # If satisfied, compile the full transcript
        transcript = None
        if satisfied:
            transcript = "=== AP Process Discovery Session ===\n\n"
            for msg in messages:
                role_label = "Client" if msg.role == "user" else "Consultant"
                transcript += f"{role_label}: {msg.content}\n\n"
            transcript += f"Consultant: {clean_response}\n\n"
            transcript += "=== End of Discovery Session ===\n"

        return {
            "status": "success",
            "data": {
                "response": clean_response,
                "satisfied": satisfied,
                "transcript": transcript,
            }
        }

    except Exception as e:
        import traceback
        traceback.print_exc()
        raise HTTPException(status_code=500, detail=f"Live chat error: {str(e)}")


class LiveChatSessionPayload(BaseModel):
    messages: List[LiveChatMessage]
    language: str = "English"


@app.post("/api/live-chat/session")
async def save_session_endpoint(payload: LiveChatSessionPayload, user_id: str = Depends(get_current_user)):
    """Save user's temporary live chat session in DB."""
    messages_list = [msg.dict() for msg in payload.messages]
    success = save_live_chat_session(user_id, messages_list, payload.language)
    if not success:
        raise HTTPException(status_code=500, detail="Failed to save live chat session to database")
    return {"status": "success"}


@app.get("/api/live-chat/session")
async def get_session_endpoint(user_id: str = Depends(get_current_user)):
    """Retrieve user's temporary live chat session from DB."""
    session_data = get_live_chat_session(user_id)
    return {"status": "success", "data": {"messages": session_data["messages"], "language": session_data["language"]}}


@app.delete("/api/live-chat/session")
async def delete_session_endpoint(user_id: str = Depends(get_current_user)):
    """Delete user's temporary live chat session from DB."""
    delete_live_chat_session(user_id)
    return {"status": "success"}


# ── Resume Prompt Cache ──────────────────────────────────────────────────────
# Simple in-memory cache: user_id -> pre-generated system prompt for resume.
# The prepare-resume endpoint writes here; the WebSocket endpoint reads + deletes.
_resume_prompt_cache: dict = {}


class PrepareResumeRequest(BaseModel):
    language: Optional[str] = None

@app.post("/api/live-chat/prepare-resume")
async def prepare_resume_endpoint(request: PrepareResumeRequest = None, user_id: str = Depends(get_current_user)):
    """
    PHASE 1 of Resume Chat: Agent 1 (Director) analyzes the saved transcript
    and generates a dynamic system prompt. The prompt is cached server-side
    so the WebSocket can pick it up instantly without blocking.
    """
    session_data = get_live_chat_session(user_id)
    saved_messages = session_data["messages"]
    
    # Use the newly requested language if provided, otherwise fallback to the saved one
    session_language = request.language if (request and request.language) else session_data.get("language", "English")
    
    # Optional: Update the session in DB with the new language so it's consistent
    if request and request.language and session_data.get("language") != request.language:
        save_live_chat_session(user_id, saved_messages, request.language)
    
    if not saved_messages:
        raise HTTPException(status_code=404, detail="No saved session found to resume")

    # Build transcript text for Agent 1
    transcript_parts = []
    for msg in saved_messages:
        role_label = "Client" if msg.get("role") == "user" else "Aria"
        content_text = (msg.get("content") or "").strip()
        if content_text:
            transcript_parts.append(f"{role_label}: {content_text}")

    history_text = "\n".join(transcript_parts)

    director_prompt = f"""You are an expert AI orchestrator.
A voice agent named Aria was interviewing a client about their Accounts Payable process. The call disconnected.
NOTE: The conversation was conducted in {session_language}.
Your output checklist must be in English (for internal AI processing reliability).
The resume session will be automatically locked to {session_language} via a separate language rule.
Here is the transcript of their conversation so far:
<transcript>
{history_text}
</transcript>

Aria's goal is to collect 8 domains:
1. PROCESS IDENTITY (Name, trigger, start/end)
2. STEP-BY-STEP WORKFLOW (Actions, who does it, inputs/outputs)
3. SYSTEMS & TOOLS (ERP, CRM, spreadsheets used)
4. PEOPLE & ROLES (Job titles, total FTE/headcount, approvers)
5. VOLUME & TIME (Frequency, annual volume, duration)
6. PAIN POINTS & ERRORS (Bottlenecks, rework, exception rate %)
7. COST & BUSINESS IMPACT (Currency, loaded hourly rate, error costs)
8. EXISTING AUTOMATION & CONSTRAINTS (Current bots, IT policies)

CRITICAL: Five facts drive the ROI maths — currency, loaded hourly rate, annual transaction volume, total FTE, and exception rate %. Any of these still missing must be the TOP PRIORITY items in the STILL MUST COLLECT section below.

Analyze the transcript and generate a STRICT two-part markdown block for Aria's new system prompt.

CRITICAL RULE ON PARTIAL DOMAINS:
If a domain is even partially answered (e.g., they gave the trigger but not the end step for Domain 1), consider that entire domain GATHERED. Put the details you learned in the FACTS section and DO NOT put that domain in the STILL MUST COLLECT section.

FORMAT EXACTLY LIKE THIS (preserve the separator lines):
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
FACTS ALREADY GATHERED (DO NOT ASK ABOUT THESE AGAIN)
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
[List exactly what the client has already stated. Be specific with names/numbers. E.g. "Systems: They use SAP and Excel."]

━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
INFORMATION YOU STILL MUST COLLECT
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
[List ONLY the domains that are completely missing. DO NOT repeat any domain listed above.]"""

    print("INFO: [prepare-resume] Triggering Agent 1 (Director) to build dynamic checklist...")
    try:
        dynamic_checklist = call_gemini_text(
            prompt=director_prompt,
            system_prompt="You are an AI Orchestrator. Output ONLY the two-part markdown checklist block, nothing else.",
            model=config.GEMINI_MODEL_FAST
        )
        safe_preview = dynamic_checklist.encode('ascii', errors='replace').decode('ascii')[:300]
        print(f"INFO: [prepare-resume] Agent 1 (Director) generated checklist (preview):\n{safe_preview}")
    except Exception as e:
        safe_err = str(e).encode('ascii', errors='replace').decode('ascii')
        print(f"ERROR: [prepare-resume] Agent 1 failed: {safe_err}")
        raise HTTPException(status_code=500, detail=f"Failed to prepare resume prompt: {safe_err}")

    # Build the full system prompt and cache it along with language
    effective_system_prompt = LIVE_CHAT_BASE_INSTRUCTIONS + "\n\n" + dynamic_checklist
    _resume_prompt_cache[user_id] = {
        "prompt": effective_system_prompt,
        "language": session_language
    }
    print(f"INFO: [prepare-resume] Cached resume prompt for user {user_id} ({len(effective_system_prompt)} chars) in {session_language}")

    return {
        "status": "success",
        "message_count": len(saved_messages)
    }


@app.websocket("/api/ws/live-chat")
async def live_chat_websocket(websocket: WebSocket):
    # Authenticate via cookie (browser automatically sends HttpOnly cookies over WS)
    # or via token query parameter (for cross-site deployments where third-party cookies are blocked)
    token = websocket.cookies.get("access_token")
    if not token:
        token = websocket.query_params.get("token")
    if not token:
        await websocket.close(code=1008)
        return
    try:
        payload = verify_token(token)
        user_id = payload.get("sub")
        if not user_id:
            await websocket.close(code=1008)
            return
    except Exception:
        await websocket.close(code=1008)
        return

    await websocket.accept()

    is_resume = websocket.query_params.get("resume", "") == "true"
    saved_messages = []

    # Read language from query param (safe default = English)
    selected_language = websocket.query_params.get("lang", "English").strip()
    # Sanitize — only accept known language names to prevent prompt injection
    ALLOWED_LANGUAGES = {
        "English", "Hindi", "Spanish", "French", "German",
        "Portuguese", "Arabic", "Japanese", "Korean", "Chinese (Simplified)"
    }
    if selected_language not in ALLOWED_LANGUAGES:
        selected_language = "English"

    # Build the language rule injection
    language_rule = f"""
LANGUAGE RULE (STRICT — NON-NEGOTIABLE):
- You MUST speak EXCLUSIVELY in {selected_language} throughout this entire conversation.
- NEVER switch to any other language, even if the user speaks to you in a different language.
- If the user speaks in a different language than {selected_language}, gently redirect them:
  respond in {selected_language} and ask them to continue in {selected_language}.
- The [DISCOVERY_COMPLETE] marker must ALWAYS be written exactly as [DISCOVERY_COMPLETE] in English, regardless of the conversation language.
"""

    # --- Dynamic System Prompt: start with the default (fresh session) ---
    effective_system_prompt = language_rule + "\n\n" + LIVE_CHAT_BASE_INSTRUCTIONS + "\n\n" + LIVE_CHAT_DEFAULT_CHECKLIST

    if is_resume:
        # PHASE 2: The frontend already called POST /api/live-chat/prepare-resume
        # which ran Agent 1 and cached the dynamic system prompt. Just pick it up.
        cached = _resume_prompt_cache.pop(user_id, None)
        if cached:
            # We explicitly trust the frontend's ?lang= query param over the cached language 
            # to allow the user to switch languages smoothly upon resume.
            
            # Re-build language rule in case it was cached differently
            language_rule = f"""
LANGUAGE RULE (STRICT — NON-NEGOTIABLE):
- You MUST speak EXCLUSIVELY in {selected_language} throughout this entire conversation.
- NEVER switch to any other language, even if the user speaks to you in a different language.
- If the user speaks in a different language than {selected_language}, gently redirect them:
  respond in {selected_language} and ask them to continue in {selected_language}.
- The [DISCOVERY_COMPLETE] marker must ALWAYS be written exactly as [DISCOVERY_COMPLETE] in English, regardless of the conversation language.
"""
            effective_system_prompt = language_rule + "\n\n" + cached.get("prompt", "")
            print(f"INFO: [ws] Using pre-cached resume prompt for user {user_id} in {selected_language}")
        else:
            print(f"WARN: [ws] No cached resume prompt found for user {user_id}, using default checklist")

        # Load saved messages so the trigger message is sent after setup
        try:
            session_data = get_live_chat_session(user_id)
            saved_messages = session_data["messages"]
            print(f"INFO: [ws] Resume mode — loaded {len(saved_messages)} messages from DB")
        except Exception as e:
            print(f"WARN: [ws] Failed to load saved messages: {e}")

    # Get API key
    api_key = config.GEMINI_LIVE_API_KEY
    if not api_key:
        await websocket.send_json({"error": "Missing Live API Key"})
        await websocket.close(code=1011)
        return

    # WebSocket connection to Gemini Live API
    url = f"wss://generativelanguage.googleapis.com/ws/google.ai.generativelanguage.v1beta.GenerativeService.BidiGenerateContent?key={api_key}"

    try:
        async with websockets.connect(url) as gemini_ws:
            # 1. Send Setup Config with the dynamically rewritten system prompt.
            # For fresh sessions: BASE_INSTRUCTIONS + DEFAULT_CHECKLIST
            # For resumed sessions: BASE_INSTRUCTIONS + DIRECTOR's custom checklist
            setup_msg = {
                "setup": {
                    "model": "models/gemini-2.5-flash-native-audio-latest",
                    "generationConfig": {
                        "responseModalities": ["AUDIO"],
                        "speechConfig": {
                            "voiceConfig": {
                                "prebuiltVoiceConfig": {
                                    "voiceName": "Aoede"
                                }
                            }
                        }
                    },
                    "systemInstruction": {
                        "parts": [{"text": effective_system_prompt}]
                    },
                    "inputAudioTranscription": {},
                    "outputAudioTranscription": {}
                }
            }
            await gemini_ws.send(json.dumps(setup_msg))

            # 2. Wait for setupComplete before starting bridge
            #    This ensures Gemini is fully ready before we relay audio
            setup_complete = False
            while not setup_complete:
                raw = await asyncio.wait_for(gemini_ws.recv(), timeout=30)
                if isinstance(raw, bytes):
                    raw_str = raw.decode("utf-8")
                else:
                    raw_str = raw
                try:
                    parsed = json.loads(raw_str)
                    if "setupComplete" in parsed:
                        setup_complete = True
                        # Forward setupComplete to frontend
                        await websocket.send_text(raw_str)
                        print("INFO: Gemini setupComplete received — bridging starts now")

                    else:
                        # Forward any other pre-setup messages (e.g. errors)
                        await websocket.send_text(raw_str)
                except Exception:
                    await websocket.send_text(raw_str)

            # 3. Send a NATURAL user trigger to kickstart the resume
            if is_resume and saved_messages:
                trigger_msg = {
                    "clientContent": {
                        "turns": [
                            {
                                "role": "user",
                                "parts": [{"text": "Hi Aria, my connection dropped but I am back online. Where were we?"}]
                            }
                        ],
                        "turnComplete": True
                    }
                }
                await gemini_ws.send(json.dumps(trigger_msg))
                print("INFO: Sent natural reconnect trigger to Aria")
            elif is_resume:
                trigger_msg = {
                    "clientContent": {
                        "turns": [
                            {
                                "role": "user",
                                "parts": [{"text": "I am back. Let's resume."}]
                            }
                        ],
                        "turnComplete": True
                    }
                }
                await gemini_ws.send(json.dumps(trigger_msg))

            # 4. Bridge the loops (now fully active)
            async def send_to_gemini():
                try:
                    while True:
                        data = await websocket.receive_text()
                        await gemini_ws.send(data)
                except WebSocketDisconnect:
                    pass
                except Exception as e:
                    print(f"Error sending to Gemini: {e}")

            async def receive_from_gemini():
                try:
                    async for message in gemini_ws:
                        if isinstance(message, bytes):
                            message = message.decode("utf-8")
                        await websocket.send_text(message)
                except Exception as e:
                    print(f"Error receiving from Gemini: {e}")

            # Bridge the loops with FIRST_COMPLETED so that if either connection drops,
            # we clean up the other connection immediately.
            done, pending = await asyncio.wait(
                [
                    asyncio.create_task(send_to_gemini()),
                    asyncio.create_task(receive_from_gemini())
                ],
                return_when=asyncio.FIRST_COMPLETED
            )
            # Cancel any pending tasks to ensure clean exit and close connections
            for task in pending:
                task.cancel()

    except asyncio.TimeoutError:
        print("ERROR: Timed out waiting for Gemini setupComplete")
        try:
            await websocket.send_json({"error": "Gemini setup timed out. Check your API key credits."})
            await websocket.close(code=1011)
        except Exception:
            pass
    except Exception as e:
        import traceback
        err_str = traceback.format_exc()
        safe_err = err_str.encode('ascii', errors='replace').decode('ascii')
        print(f"WS CRASH: {safe_err}")
        try:
            safe_msg = str(e).encode('ascii', errors='replace').decode('ascii')
            await websocket.send_json({"error": safe_msg})
            await websocket.close(code=1011)
        except Exception:
            pass




# ============================================================
# Public Developer API (v1) -- named models via API key (ap_sk_...)
# ============================================================
V1_MODELS = {"ap_analysis", "ap_pdd-sdd", "ap_bpmn"}


def _build_analysis_output(result: dict) -> dict:
    """Curate the pipeline result into the stateless analysis payload.

    This is the SAME shape consumed by ap_pdd-sdd / ap_bpmn, so callers can
    chain models (analysis -> document/bpmn) without losing fidelity.
    """
    keys = [
        "original_filename", "input_format", "process_map", "scored_steps",
        "priority_targets", "opportunities", "executive_summary",
        "markdown_report", "systems_mentioned", "roles_identified",
        "pain_points", "hourly_rate", "project_timeline", "roi_estimate",
        "missing_critical_facts",
    ]
    return {k: result.get(k) for k in keys if result.get(k) is not None}


def _run_pipeline_analysis(inp: dict) -> dict:
    """ap_analysis: stateless run accepting EITHER 'text' OR base64 'file_b64' + 'filename'
    (audio / video / docx / txt). Mirrors the /api/analyze pre-processing exactly.
    """
    text = (inp.get("text") or "").strip()
    file_b64 = inp.get("file_b64") or inp.get("file_base64")
    filename = inp.get("filename") or ""

    if file_b64:
        ext = os.path.splitext(filename)[1].lower()
        if ext not in ALL_SUPPORTED_EXTENSIONS:
            raise HTTPException(status_code=415, detail=f"Unsupported file type '{ext}'. Supported: {sorted(ALL_SUPPORTED_EXTENSIONS)}")
        try:
            contents = base64.b64decode(file_b64)
        except Exception:
            raise HTTPException(status_code=400, detail="'file_b64' is not valid base64.")
        size_mb = len(contents) / (1024 * 1024)
        if size_mb > MAX_UPLOAD_SIZE_MB:
            raise HTTPException(status_code=413, detail=f"File too large ({size_mb:.1f}MB). Max: {MAX_UPLOAD_SIZE_MB}MB")
        temp_filename = f"{uuid.uuid4()}{ext}"
        temp_path = os.path.join(TEMP_DIR, temp_filename)
        with open(temp_path, "wb") as f:
            f.write(contents)
        input_format = detect_input_type(filename)
        try:
            result = pipeline.invoke({
                "raw_text": "",
                "input_format": input_format,
                "original_filename": filename,
                "file_path": temp_path,
            })
        finally:
            try:
                os.remove(temp_path)
            except Exception:
                pass
    elif text:
        if len(text) < 50:
            raise HTTPException(status_code=400, detail="Text too short. Provide at least 50 characters.")
        result = pipeline.invoke({
            "raw_text": text,
            "input_format": "text",
            "original_filename": "",
            "file_path": "",
        })
    else:
        raise HTTPException(status_code=400, detail="Provide 'text' or 'file_b64' + 'filename' in 'input'.")

    return _build_analysis_output(result)


class V1RunRequest(BaseModel):
    model: str
    input: Optional[dict] = None


@app.post("/v1/run")
async def v1_run(payload: V1RunRequest, account: dict = Depends(get_api_account)):
    """Single entrypoint for the public API. Authenticated by an ap_sk_ key.

    Body: {"model": "ap_analysis|ap_pdd-sdd|ap_bpmn", "input": {...}}

    The named-model wrappers reuse the EXACT in-app orchestration, so internal
    agent-to-agent calls (PDD/SDD -> BPMN agent -> SVG flowchart agent) behave
    identically to the app and never break.
    """
    model = (payload.model or "").strip()
    inp = payload.input or {}
    if not isinstance(inp, dict):
        raise HTTPException(status_code=400, detail="'input' must be an object.")
    if model not in V1_MODELS:
        raise HTTPException(status_code=400, detail=f"Unknown model '{model}'. Available: {sorted(V1_MODELS)}")

    allowed = account.get("allowed_models") or []
    if allowed and model not in allowed:
        raise HTTPException(status_code=403, detail=f"This API key may not use '{model}'. Allowed: {allowed}")

    user_id = account["user_id"]

    try:
        if model == "ap_analysis":
            return {"model": model, "data": _run_pipeline_analysis(inp)}

        if model == "ap_bpmn":
            scored_steps = inp.get("scored_steps")
            if not scored_steps and isinstance(inp.get("analysis"), dict):
                scored_steps = inp["analysis"].get("scored_steps")
            if not scored_steps and (inp.get("text") or inp.get("file_b64") or inp.get("file_base64")):
                # Fallback: derive steps by running the analysis pipeline on raw text/file.
                scored_steps = _run_pipeline_analysis(inp).get("scored_steps")
            if not scored_steps:
                raise HTTPException(status_code=400, detail="Provide 'scored_steps', an 'analysis' object, or raw 'text'/'file_b64' to derive them.")
            bpmn_xml = _generate_bpmn_xml(scored_steps)
            return {"model": model, "data": {"bpmn_xml": bpmn_xml}}

        if model == "ap_pdd-sdd":
            doc_type = (inp.get("doc_type") or "pdd").lower()
            if doc_type not in ("pdd", "sdd"):
                raise HTTPException(status_code=400, detail="'doc_type' must be 'pdd' or 'sdd'.")
            analysis = inp.get("analysis") or inp.get("assessment")
            if (not isinstance(analysis, dict) or not analysis) and (inp.get("text") or inp.get("file_b64") or inp.get("file_base64")):
                # Fallback: derive the full analysis from raw text/file.
                analysis = _run_pipeline_analysis(inp)
            if not isinstance(analysis, dict) or not analysis:
                raise HTTPException(status_code=400, detail="Provide an 'analysis' object (output of ap_analysis), or raw 'text'/'file_b64' to derive it.")
            # Reuse the EXACT in-app document orchestration by materialising a
            # short-lived assessment, generating against it, then deleting it.
            # This guarantees byte-identical BPMN + SVG diagram behaviour.
            temp_data = dict(analysis)
            temp_data["original_filename"] = "[API] " + str(analysis.get("original_filename") or "document")
            temp_data["_api_temp"] = True
            temp_id = save_assessment(temp_data, user_id)
            if not temp_id:
                raise HTTPException(status_code=500, detail="Database unavailable for document generation.")
            try:
                doc_result = await generate_document(
                    temp_id,
                    DocumentRequest(doc_type=doc_type, force=True),
                    user_id=user_id,
                )
            finally:
                try:
                    delete_assessment(temp_id, user_id)
                except Exception:
                    pass
            content = doc_result.get("content", "") if isinstance(doc_result, dict) else ""
            return {"model": model, "data": {"doc_type": doc_type, "content": content}}

    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))
