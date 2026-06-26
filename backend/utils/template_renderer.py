"""
Template-driven PDD / SDD rendering.

The feature: a user uploads their own PDD/SDD Word template (with the
Auxiliobits logo, header/footer, fonts and section structure). When they
generate + download a document, the output must follow that template's design
AND structure.

HOW IT WORKS
------------
1. build_template_outline()  - reads the uploaded .docx and returns its exact
   heading structure. This is fed into the generation prompt so the generated
   markdown mirrors the template's sections, numbering and hierarchy.

2. render_markdown_into_template() - the design half. Instead of building a
   blank Word document, we open a COPY of the uploaded template, keep its
   "cover" (everything before the first heading -> this is where the logo and
   title block live), keep its header/footer and styles, drop the template's
   empty body sections, then append the generated content rendered with the
   template's OWN named styles (Heading 1/2/3, Normal, tables). Because we reuse
   the original file as the skeleton, the logo, footer and fonts come along for
   free - we never recreate them.

Diagrams (BPMN + flowchart SVGs) are embedded from the PNG snapshots the
frontend captures from the rendered DOM, exactly like the legacy download path.
"""
import io
import re
import base64

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


# ── Template structure extraction ───────────────────────────────────────────

def analyze_template(template_path: str) -> list[dict]:
    """Return the heading structure of the template as a list of
    {"text": str, "level": int} in document order."""
    doc = Document(template_path)
    headings = []
    for p in doc.paragraphs:
        style = (p.style.name if p.style else "") or ""
        if style.lower().startswith("heading") and p.text.strip():
            try:
                level = int(style.split()[-1])
            except (ValueError, IndexError):
                level = 1
            headings.append({"text": p.text.strip(), "level": level})
    return headings


def build_template_outline(template_path: str) -> str:
    """Render the template's heading structure as an indented outline string
    suitable for injecting into the generation prompt."""
    headings = analyze_template(template_path)
    lines = []
    for h in headings:
        indent = "    " * (max(h["level"], 1) - 1)
        lines.append(f"{indent}{h['text']}")
    return "\n".join(lines)


# ── Skeleton preparation ────────────────────────────────────────────────────

def _first_heading_index(body, children: list) -> int | None:
    for i, el in enumerate(children):
        if el.tag != qn("w:p"):
            continue
        pPr = el.find(qn("w:pPr"))
        if pPr is None:
            continue
        pStyle = pPr.find(qn("w:pStyle"))
        if pStyle is not None:
            val = (pStyle.get(qn("w:val")) or "").lower()
            if val.startswith("heading"):
                return i
    return None


def _trim_body_keep_cover(doc: Document) -> bool:
    """Remove the template's body content from the first heading onward, keeping
    everything before it (the cover/logo/title block) and the trailing section
    properties (which reference the header/footer). Returns True if a heading
    boundary was found."""
    body = doc.element.body
    children = list(body)
    start = _first_heading_index(body, children)
    if start is None:
        return False
    for el in children[start:]:
        if el.tag == qn("w:sectPr"):
            continue  # keep section/page setup + header/footer references
        body.remove(el)
    return True


def _ensure_sectpr_last(doc: Document) -> None:
    """Guarantee the body-level <w:sectPr> stays the final child so all appended
    content renders before it (and the footer/header stay attached)."""
    body = doc.element.body
    sectPr = body.find(qn("w:sectPr"))
    if sectPr is not None:
        body.remove(sectPr)
        body.append(sectPr)


# ── Inline + block helpers ──────────────────────────────────────────────────

def _add_runs(paragraph, text: str) -> None:
    """Add text to a paragraph, honouring **bold** markers and stripping stray
    inline markdown."""
    text = re.sub(r"`([^`]*)`", r"\1", text)  # drop inline code ticks
    for part in re.split(r"(\*\*.*?\*\*)", text):
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        else:
            paragraph.add_run(part)


def _pick_style(doc: Document, candidates: list[str]) -> str | None:
    for name in candidates:
        try:
            doc.styles[name]
            return name
        except KeyError:
            continue
    return None


def _fit_table(tbl) -> None:
    tbl.autofit = True
    tblPr = tbl._tbl.tblPr
    for el in tblPr.findall(qn("w:tblW")):
        tblPr.remove(el)
    tblW = OxmlElement("w:tblW")
    tblW.set(qn("w:w"), "5000")
    tblW.set(qn("w:type"), "pct")
    tblPr.append(tblW)
    for row in tbl.rows:
        for cell in row.cells:
            for para in cell.paragraphs:
                for run in para.runs:
                    run.font.size = Pt(8.5)


class _ImagePlacer:
    """Embeds frontend-captured diagram PNGs, in document order."""

    def __init__(self, doc: Document, images: list, max_w, max_h):
        self.doc = doc
        self.max_w = max_w
        self.max_h = max_h
        self.bpmn = [im for im in images if im.get("type") == "bpmn" and im.get("png_base64")]
        self.svg = [im for im in images if im.get("type") == "svg" and im.get("png_base64")]
        self.bpmn_i = 0
        self.svg_i = 0

    def _embed_b64(self, png_b64: str) -> bool:
        try:
            if "," in png_b64:
                png_b64 = png_b64.split(",", 1)[1]
            data = base64.b64decode(png_b64)
            self._add_fitted(data)
            return True
        except Exception as e:
            print(f"[template_renderer] image embed failed: {e}")
            return False

    def _add_fitted(self, data: bytes) -> None:
        pw = ph = 0
        try:
            from PIL import Image as _PILImage
            with _PILImage.open(io.BytesIO(data)) as im:
                pw, ph = im.size
        except Exception:
            pass
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        if pw and ph and int(self.max_w * (ph / pw)) <= self.max_h:
            run.add_picture(io.BytesIO(data), width=self.max_w)
        elif pw and ph:
            run.add_picture(io.BytesIO(data), height=self.max_h)
        else:
            run.add_picture(io.BytesIO(data), width=self.max_w)

    def place_bpmn(self) -> None:
        if self.bpmn_i < len(self.bpmn):
            self._embed_b64(self.bpmn[self.bpmn_i].get("png_base64", ""))
            self.bpmn_i += 1
        else:
            self.doc.add_paragraph("[Process diagram not captured]")

    def place_svg(self) -> None:
        if self.svg_i < len(self.svg):
            self._embed_b64(self.svg[self.svg_i].get("png_base64", ""))
            self.svg_i += 1
        else:
            self.doc.add_paragraph("[Diagram not captured \u2014 re-open and try again]")


# ── Markdown -> docx body ───────────────────────────────────────────────────

def _render_markdown(doc: Document, content: str, images: list) -> None:
    section = doc.sections[0]
    content_w = section.page_width - section.left_margin - section.right_margin
    content_h = section.page_height - section.top_margin - section.bottom_margin
    max_img_w = min(content_w, Inches(6.5))
    max_img_h = int(content_h * 0.9)
    placer = _ImagePlacer(doc, images, max_img_w, max_img_h)

    table_style = _pick_style(doc, ["Table Grid", "Light Grid Accent 1", "Table Normal"])
    bullet_style = _pick_style(doc, ["List Bullet", "List Paragraph"])

    # Flatten <figure class="doc-flowchart">...</figure> (multi-line SVG) into a
    # single sentinel BEFORE tokenising by newline.
    content = re.sub(
        r'<figure[^>]*class=["\']doc-flowchart["\'][^>]*>[\s\S]*?</figure>',
        "\n[[FLOWCHART_FIGURE]]\n",
        content,
        flags=re.MULTILINE,
    )

    lines = content.split("\n")
    i = 0
    table_rows: list[list[str]] = []
    in_code = False
    code_lang = ""
    code_lines: list[str] = []

    def flush_table():
        nonlocal table_rows
        if not table_rows:
            return
        cols = max(len(r) for r in table_rows)
        tbl = doc.add_table(rows=len(table_rows), cols=cols)
        if table_style:
            try:
                tbl.style = table_style
            except KeyError:
                pass
        for ri, row in enumerate(table_rows):
            for ci in range(cols):
                cell = tbl.rows[ri].cells[ci]
                cell.text = ""
                value = row[ci] if ci < len(row) else ""
                para = cell.paragraphs[0]
                _add_runs(para, value)
                if ri == 0:
                    for run in para.runs:
                        run.bold = True
        _fit_table(tbl)
        table_rows = []

    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        # Code fences (bpmn / svg / generic)
        if stripped.startswith("```"):
            if not in_code:
                in_code = True
                code_lang = stripped[3:].strip().lower()
                code_lines = []
            else:
                in_code = False
                flush_table()
                if code_lang == "bpmn":
                    placer.place_bpmn()
                elif code_lang in ("svg", "xml", "html") and "<svg" in "\n".join(code_lines):
                    placer.place_svg()
                else:
                    p = doc.add_paragraph()
                    run = p.add_run("\n".join(code_lines))
                    run.font.name = "Consolas"
                    run.font.size = Pt(9)
                    run.font.color.rgb = RGBColor(0x60, 0x60, 0x60)
            i += 1
            continue
        if in_code:
            code_lines.append(line)
            i += 1
            continue

        # Tables
        if stripped.startswith("|") and "|" in stripped:
            cells = [c.strip() for c in stripped.strip("|").split("|")]
            if all(set(c).issubset({"-", ":", " "}) for c in cells):
                i += 1
                continue
            table_rows.append(cells)
            i += 1
            continue
        else:
            flush_table()

        # Headings -> reuse the template's own Heading styles
        if stripped.startswith("#### "):
            doc.add_heading(stripped[5:], level=3)
        elif stripped.startswith("### "):
            doc.add_heading(stripped[4:], level=3)
        elif stripped.startswith("## "):
            doc.add_heading(stripped[3:], level=2)
        elif stripped.startswith("# "):
            doc.add_heading(stripped[2:], level=1)
        elif stripped == "[[FLOWCHART_FIGURE]]":
            placer.place_svg()
        elif stripped in ("[[BPMN_DIAGRAM]]", "{BPMN_DIAGRAM}"):
            placer.place_bpmn()
        elif stripped.startswith("---"):
            pass  # horizontal rule -> skip (template provides its own spacing)
        elif stripped.startswith(("- ", "* ")):
            p = doc.add_paragraph(style=bullet_style) if bullet_style else doc.add_paragraph()
            _add_runs(p, stripped[2:])
        elif stripped:
            p = doc.add_paragraph()
            _add_runs(p, stripped)
        i += 1

    flush_table()


# ── Public entry point ──────────────────────────────────────────────────────

def render_markdown_into_template(template_path: str, content: str, images: list | None = None) -> bytes:
    """Render the generated markdown into a copy of the uploaded template,
    preserving the template's logo, header/footer, fonts and cover. Returns the
    resulting .docx as bytes."""
    images = images or []
    doc = Document(template_path)
    _trim_body_keep_cover(doc)
    _render_markdown(doc, content, images)
    _ensure_sectpr_last(doc)
    out = io.BytesIO()
    doc.save(out)
    return out.getvalue()
